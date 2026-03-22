"""
Generiert die Kapitel-3-Dokumentation als Word-Datei (.docx)
Ausführen: python generate_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Seitenränder ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ─── Hilfsfunktionen ─────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, italic=False, color=None, name="Calibri"):
    run.bold   = bold
    run.italic = italic
    run.font.name = name
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 6)
    p.paragraph_format.space_after  = Pt(4)
    return p

def para(text="", bold=False, size=11, space_after=6, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, color=color)
    return p

def code_block(lines):
    """Grauer Code-Block mit Monospace-Schrift"""
    if isinstance(lines, str):
        lines = [lines]
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        # Grauer Hintergrund via paragraph shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F2F2F2")
        pPr.append(shd)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    # Abstand nach dem Block
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def info_box(text, color_fill="E8F4FD", color_border="2196F3"):
    """Blauer Info-Kasten"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_fill)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"

def warning_box(text):
    """Oranger Warnkasten für Live-Demo"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "FFF3E0")
    pPr.append(shd)
    run = p.add_run("⚠  " + text)
    run.font.size = Pt(10)
    run.bold = True
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)

def demo_step(number, title, commands=None, note=None):
    """Nummerierter Demo-Schritt"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"Schritt {number} – {title}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    if note:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(f"→  {note}")
        r2.italic = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    if commands:
        code_block(commands)

def simple_table(headers, rows, col_widths=None):
    """Erstellt eine formatierte Tabelle"""
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "1A73E8")
        tcPr.append(shd)
    # Zeilen
    for ri, row_data in enumerate(rows):
        row = t.rows[ri+1]
        fill = "FFFFFF" if ri % 2 == 0 else "F8F9FA"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  fill)
            tcPr.append(shd)
    # Spaltenbreiten
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t

def divider():
    p = doc.add_paragraph("─" * 80)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(8)

def live_demo_banner(section_num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "D32F2F")
    pPr.append(shd)
    r = p.add_run(f"  LIVE-DEMO {section_num}  –  {title}  |  Bildschirmaufnahme erforderlich!")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.name = "Calibri"

# ═════════════════════════════════════════════════════════════════════════════
# TITELSEITE
# ═════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph("ElectroSwap")
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.runs[0]
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

t2 = doc.add_paragraph("Premium Hardware E-Commerce")
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.runs[0]
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
t3 = doc.add_paragraph("Kapitel 3 – Datenbankoperationen und Architektur")
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.runs[0]
r3.bold = True
r3.font.size = Pt(16)

doc.add_paragraph()
meta = doc.add_paragraph("Stack: Python / Flask  ·  MongoDB 7.0  ·  Nginx  ·  Docker\nDatum: 2026-03-21")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in meta.runs:
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 3.1 – KONZEPT ZUGRIFFSBERECHTIGUNGEN (DOKU)
# ═════════════════════════════════════════════════════════════════════════════

heading("3.1  Konzept für Zugriffsberechtigungen", level=1)

info_box(
    "Aufgabe: Konzept für Zugriffsberechtigungen (Rollen, Benutzer, Berechtigungen) "
    "und alle relevanten Befehle in der Dokumentation."
)

heading("Zwei-Ebenen-Modell", level=2)

para(
    "ElectroSwap sichert den Datenzugriff auf zwei unabhängigen Ebenen. "
    "Beide Ebenen arbeiten ergänzend: Die Applikationsebene schützt die Web-Routen, "
    "die Datenbankebene schützt MongoDB vor direktem Zugriff."
)

code_block([
    "┌─────────────────────────────────────────────────────┐",
    "│           Ebene 1: Applikation (Flask)              │",
    "│   Rollen: customer / admin                          │",
    "│   Schutz: @login_required + @admin_required         │",
    "├─────────────────────────────────────────────────────┤",
    "│           Ebene 2: Datenbank (MongoDB)              │",
    "│   Benutzer: readonly / app / admin / root           │",
    "│   Auth:    SCRAM-SHA-256                            │",
    "└─────────────────────────────────────────────────────┘",
])

# ── Applikationsebene ──────────────────────────────────────────────────────

heading("Ebene 1: Applikationsrollen (Flask RBAC)", level=2)

para("Jeder registrierte Benutzer erhält im MongoDB-Dokument eine Rolle. "
     "Diese Rolle steuert, welche Routen zugänglich sind.")

simple_table(
    ["Rolle", "Vergabe", "Zugang"],
    [
        ["customer", "Automatisch bei Registrierung", "Produkte, Warenkorb, Bestellungen, Bewertungen"],
        ["admin",    "Manuell durch DB-Admin",         "Alles + Admin-Dashboard, Produkt-CRUD, Bestellstatus"],
    ],
    col_widths=[3.5, 5.5, 7.5]
)

heading("Berechtigungsmatrix (Applikation)", level=3)

simple_table(
    ["Funktion", "customer", "admin"],
    [
        ["Produkte ansehen / suchen / filtern", "✓", "✓"],
        ["Warenkorb & Wunschliste",             "✓", "✓"],
        ["Bestellung aufgeben",                 "✓", "✓"],
        ["Bestellhistorie ansehen",             "✓", "✓"],
        ["Produktbewertungen schreiben",        "✓", "✓"],
        ["Profil bearbeiten",                   "✓", "✓"],
        ["Admin-Dashboard",                     "✗", "✓"],
        ["Produkte erstellen / bearbeiten / löschen", "✗", "✓"],
        ["Bestellstatus ändern",                "✗", "✓"],
    ],
    col_widths=[10, 3, 3.5]
)

heading("Speicherung der Rolle im User-Dokument", level=3)

code_block([
    '{',
    '  "_id":           "ObjectId(...)",',
    '  "username":      "max",',
    '  "email":         "max@example.com",',
    '  "password_hash": "$2b$12$...",',
    '  "role":          "customer",',
    '  "address":       { ... },',
    '  "created_at":    "ISODate(...)"',
    '}',
])

heading("Schutz der Admin-Routen (Code)", level=3)

para("Der Decorator @admin_required in app/admin/routes.py schützt alle Admin-Endpunkte:")

code_block([
    "def admin_required(f):",
    "    @wraps(f)",
    "    @login_required",
    "    def wrapper(*args, **kwargs):",
    "        if not current_user.is_admin:",
    "            flash('Access denied.', 'error')",
    "            return redirect(url_for('main.index'))",
    "        return f(*args, **kwargs)",
    "    return wrapper",
])

heading("Geschützte Admin-Routen", level=3)

code_block([
    "GET/POST  /admin/                      → Dashboard",
    "GET       /admin/products              → Produktliste",
    "POST      /admin/products/create       → Produkt erstellen",
    "GET/POST  /admin/products/edit/<id>    → Produkt bearbeiten",
    "POST      /admin/products/delete/<id>  → Produkt löschen",
    "GET       /admin/orders                → Bestellübersicht",
    "POST      /admin/orders/<id>/status    → Status ändern",
])

heading("Passwort-Sicherheit", level=3)

para("Passwörter werden mit bcrypt (Salt + Hash) gespeichert – niemals im Klartext:")

code_block([
    "# Registrierung – Passwort hashen",
    'pw_hash = bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")',
    "",
    "# Login – Passwort prüfen",
    'bcrypt.checkpw(password.encode("utf-8"), stored_hash.encode("utf-8"))',
])

# ── Datenbankebene ─────────────────────────────────────────────────────────

heading("Ebene 2: Datenbankbenutzer (MongoDB)", level=2)

para("MongoDB unterstützt eigene Benutzer mit feingranularen Rollen. "
     "ElectroSwap verwendet vier Datenbankbenutzer nach dem Least-Privilege-Prinzip "
     "(jeder Benutzer hat nur die minimal nötigen Rechte).")

simple_table(
    ["DB-Benutzer", "MongoDB-Rolle", "Passwort", "Verwendung"],
    [
        ["root",                  "root (admin-DB)", "RootPass000!",  "MongoDB-Superuser – nur für Setup"],
        ["electroswap_admin",     "dbOwner",         "AdminPass123!", "Schema-Änderungen, Restore"],
        ["electroswap_app",       "readWrite",        "AppPass456!",  "Flask-Applikation (lesen + schreiben)"],
        ["electroswap_readonly",  "read",             "ReadPass789!", "Backup und Reporting (nur lesen)"],
    ],
    col_widths=[4.5, 4, 3.5, 5.5]
)

heading("Berechtigungsmatrix (Datenbank)", level=3)

simple_table(
    ["Operation", "root", "admin", "app", "readonly"],
    [
        ["Lesen (find)",                     "✓", "✓", "✓", "✓"],
        ["Schreiben (insert/update/delete)",  "✓", "✓", "✓", "✗"],
        ["Collections erstellen/löschen",    "✓", "✓", "✗", "✗"],
        ["Indexes erstellen",                "✓", "✓", "✗", "✗"],
        ["Backup (mongodump)",               "✓", "✓", "✗", "✓"],
        ["Restore (mongorestore --drop)",    "✓", "✓", "✗", "✗"],
        ["Datenbankbenutzer verwalten",      "✓", "✓", "✗", "✗"],
    ],
    col_widths=[7, 1.5, 1.8, 1.5, 2.5]
)

heading("Automatische Erstellung (mongo-init/init-users.js)", level=3)

para("Die Benutzer werden beim ersten Container-Start automatisch angelegt:")

code_block([
    "// electroswap_admin (dbOwner – Vollzugriff)",
    'db.getSiblingDB("electroswap").createUser({',
    '  user: "electroswap_admin",',
    '  pwd:  "AdminPass123!",',
    '  roles: [{ role: "dbOwner", db: "electroswap" }]',
    "});",
    "",
    "// electroswap_app (readWrite – Flask-Applikation)",
    'db.getSiblingDB("electroswap").createUser({',
    '  user: "electroswap_app",',
    '  pwd:  "AppPass456!",',
    '  roles: [{ role: "readWrite", db: "electroswap" }]',
    "});",
    "",
    "// electroswap_readonly (read – Backup/Reporting)",
    'db.getSiblingDB("electroswap").createUser({',
    '  user: "electroswap_readonly",',
    '  pwd:  "ReadPass789!",',
    '  roles: [{ role: "read", db: "electroswap" }]',
    "});",
])

heading("Verbindungsbefehle (mongosh) – getestet & funktionierend", level=3)

para(
    "Hinweis: Das '!' im Passwort muss in Verbindungs-URLs als '%21' kodiert werden. "
    "Die --password-Flag funktioniert mit dem Literal '!' direkt."
)

code_block([
    "# Format: mongosh <db> --username X --password Y --authenticationDatabase Z",
    "",
    "# Als electroswap_admin (dbOwner – Vollzugriff)",
    "docker exec -it electroswap_mongo mongosh electroswap \\",
    '  --username electroswap_admin \\',
    '  --password "AdminPass123!" \\',
    "  --authenticationDatabase electroswap",
    "",
    "# Als electroswap_app (readWrite – Flask-App-User)",
    "docker exec -it electroswap_mongo mongosh electroswap \\",
    '  --username electroswap_app \\',
    '  --password "AppPass456!" \\',
    "  --authenticationDatabase electroswap",
    "",
    "# Als electroswap_readonly (read – nur Lesen)",
    "docker exec -it electroswap_mongo mongosh electroswap \\",
    '  --username electroswap_readonly \\',
    '  --password "ReadPass789!" \\',
    "  --authenticationDatabase electroswap",
])

heading("Verwaltungsbefehle (mongosh)", level=3)

code_block([
    "// Alle MongoDB-Datenbankbenutzer anzeigen",
    "use electroswap",
    "db.getUsers()",
    "",
    "// Applikations-User (customer/admin) anzeigen",
    "db.users.find({}, { username: 1, email: 1, role: 1, _id: 0 })",
    "",
    "// Nur Admins anzeigen",
    'db.users.find({ role: "admin" }, { username: 1, email: 1, _id: 0 })',
    "",
    "// Benutzer zum Admin befördern",
    "db.users.updateOne(",
    '  { email: "max@example.com" },',
    '  { $set: { role: "admin" } }',
    ")",
    "",
    "// Wieder auf customer zurücksetzen",
    "db.users.updateOne(",
    '  { email: "max@example.com" },',
    '  { $set: { role: "customer" } }',
    ")",
])

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 3.2 – LIVE-DEMO ZUGRIFFSBERECHTIGUNGEN
# ═════════════════════════════════════════════════════════════════════════════

live_demo_banner("3.2", "Zugriffsberechtigungen zeigen – Rollen / Benutzer")

warning_box(
    "BILDSCHIRMAUFNAHME STARTEN! Terminal und Browser nebeneinander sichtbar halten. "
    "Alle 4 Blöcke zeigen – nichts überspringen. Demo-Skript pausiert automatisch zwischen den Blöcken."
)

info_box(
    "Ziel: A) customer vs. admin im Browser zeigen  B) 3 MongoDB-User mit getUsers()  "
    "C) Readonly kann nicht schreiben (Fehler zeigen)  D) Rolle live ändern → Browser-Effekt zeigen"
)

heading("Vorbereitung (vor Aufnahme starten)", level=2)

code_block([
    "# 1. Stack starten",
    "cd C:\\Projects\\ElectroSwap\\ElectroSwap",
    "docker compose up --build -d",
    "",
    "# 2. Alle Services prüfen",
    "docker compose ps",
    "# → 5 Container müssen Up/healthy sein",
    "",
    "# 3. Seed-Daten sicherstellen",
    "docker exec electroswap_app1 python seed_data.py",
    "",
    "# 4. Browser öffnen: http://localhost → Startseite prüfen",
    "# 5. Terminal im Projektverzeichnis bereithalten",
])

heading("Option A – Demo-Skript (empfohlen)", level=2)

para(
    "Das Skript demo_rollen.sh führt alle Schritte durch und pausiert zwischen den "
    "Blöcken, damit die Kamera alles erfassen kann:"
)

code_block([
    "bash demo_rollen.sh",
    "",
    "# Das Skript zeigt automatisch:",
    "#   A) Alle App-User mit Rollen (find + aggregate)",
    "#   B) db.getUsers() → 3 MongoDB-Benutzer mit Rollen",
    "#   C) Readonly: insert schlägt fehl, deleteMany schlägt fehl",
    "#   D) demo_customer → admin → Rollenreset",
])

heading("Option B – Manuelle Befehle (einzeln ausführen)", level=2)

demo_step(1, "Alle Applikations-Benutzer anzeigen",
    note="Im Terminal zeigen: Wer ist customer, wer ist admin",
    commands=[
        "docker exec electroswap_mongo mongosh electroswap \\",
        '  --username electroswap_admin --password "AdminPass123!" \\',
        "  --authenticationDatabase electroswap \\",
        '  --eval \'db.users.find({},{username:1,email:1,role:1,_id:0}).toArray()\'',
        "",
        "# Ausgabe zeigt: admin (role:admin), demo_customer (role:customer), ...",
    ]
)

demo_step(2, "Als Customer im Browser – Admin ist gesperrt",
    note="Browser zeigen: customer@electroswap.ch hat KEINEN Admin-Zugriff",
    commands=[
        "# Browser: http://localhost/auth/login",
        "#   E-Mail:   customer@electroswap.ch",
        "#   Passwort: customer123",
        "#",
        "# Nach Login zeigen:",
        "#   → Kein Admin-Link im Menü vorhanden",
        "#   → http://localhost/admin/ direkt aufrufen",
        "#   → Flash-Message: 'Access denied'  (Weiterleitung zur Startseite)",
    ]
)

demo_step(3, "Als Admin im Browser – voller Zugriff",
    note="Browser: admin@electroswap.ch hat das Admin-Dashboard",
    commands=[
        "# Ausloggen, dann neu einloggen:",
        "#   E-Mail:   admin@electroswap.ch",
        "#   Passwort: admin123",
        "#",
        "# Zeigen:",
        "#   → Admin-Link im Menü erscheint",
        "#   → http://localhost/admin/  → Dashboard öffnet",
        "#   → /admin/products/         → Produkt-Tabelle + Bearbeiten-Buttons",
        "#   → /admin/orders/           → Bestellübersicht",
    ]
)

demo_step(4, "MongoDB-Datenbankbenutzer anzeigen (db.getUsers)",
    note="Terminal: 3 DB-User mit unterschiedlichen Rollen zeigen",
    commands=[
        "docker exec electroswap_mongo mongosh electroswap \\",
        '  --username electroswap_admin --password "AdminPass123!" \\',
        "  --authenticationDatabase electroswap \\",
        "  --eval 'db.getUsers()'",
        "",
        "# Ausgabe zeigt:",
        "#   electroswap_admin   → role: dbOwner",
        "#   electroswap_app     → role: readWrite",
        "#   electroswap_readonly → role: read",
    ]
)

demo_step(5, "Least-Privilege: Readonly darf NICHT schreiben (Fehler zeigen)",
    note="Terminal: Zwei Fehler zeigen – insert und deleteMany scheitern",
    commands=[
        "# Versuch 1: INSERT schlägt fehl",
        "docker exec electroswap_mongo mongosh electroswap \\",
        '  --username electroswap_readonly --password "ReadPass789!" \\',
        "  --authenticationDatabase electroswap \\",
        '  --eval \'db.users.insertOne({test:"unauthorized"})\'',
        "# → MongoServerError: not authorized ... { insert: ... }",
        "",
        "# Versuch 2: DELETE schlägt fehl",
        "docker exec electroswap_mongo mongosh electroswap \\",
        '  --username electroswap_readonly --password "ReadPass789!" \\',
        "  --authenticationDatabase electroswap \\",
        "  --eval 'db.products.deleteMany({})'",
        "# → MongoServerError: not authorized ... { delete: ... }",
    ]
)

demo_step(6, "Live-Rollenänderung: customer → admin → customer",
    note="Terminal + Browser: Rollenänderung per DB-Befehl wirkt sofort nach Login",
    commands=[
        "# Schritt 1: demo_customer zu admin befördern",
        "docker exec electroswap_mongo mongosh electroswap \\",
        '  --username electroswap_admin --password "AdminPass123!" \\',
        "  --authenticationDatabase electroswap \\",
        '  --eval \'db.users.updateOne({username:"demo_customer"},{$set:{role:"admin"}});',
        '          db.users.findOne({username:"demo_customer"},{username:1,role:1,_id:0})\'',
        "# → { username: 'demo_customer', role: 'admin' }",
        "",
        "# Schritt 2: Browser – customer@electroswap.ch neu einloggen",
        "#   → Admin-Bereich ist jetzt zugänglich!",
        "",
        "# Schritt 3: Rolle zurücksetzen",
        "docker exec electroswap_mongo mongosh electroswap \\",
        '  --username electroswap_admin --password "AdminPass123!" \\',
        "  --authenticationDatabase electroswap \\",
        '  --eval \'db.users.updateOne({username:"demo_customer"},{$set:{role:"customer"}});',
        '          db.users.findOne({username:"demo_customer"},{username:1,role:1,_id:0})\'',
        "# → { username: 'demo_customer', role: 'customer' }",
    ]
)

warning_box("Bildschirmaufnahme STOPPEN. Alle 6 Schritte + 4 Demo-Blöcke gezeigt.")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 3.3 – BACKUP-KONZEPT (DOKU)
# ═════════════════════════════════════════════════════════════════════════════

heading("3.3  Backup-Konzept", level=1)

info_box("Aufgabe: Backup-Konzept dokumentieren, alle Backup- und Restore-Befehle in Dokumentation.")

heading("Strategie", level=2)

simple_table(
    ["Aspekt", "Entscheidung", "Begründung"],
    [
        ["Tool",             "mongodump / mongorestore",     "Offizielle MongoDB-Tools, BSON-Format"],
        ["Backup-Benutzer",  "electroswap_readonly",         "Least-Privilege: Backup braucht nur Lesezugriff"],
        ["Restore-Benutzer", "electroswap_admin",            "Vollzugriff für --drop erforderlich"],
        ["Format",           "Komprimiertes Archiv (.gz)",   "Platzsparend, portabel"],
        ["Namensschema",     "electroswap_YYYYMMDD_HHMMSS",  "Eindeutig, chronologisch sortierbar"],
        ["Aufbewahrung",     "7 Tage (automatische Bereinigung)", "Ausgewogenes Verhältnis Speicher/Sicherheit"],
    ],
    col_widths=[4, 5, 7.5]
)

heading("Gesicherte Collections", level=2)

simple_table(
    ["Collection", "Inhalt"],
    [
        ["users",     "Benutzerkonten, Rollen, Passwort-Hashes, Adressen"],
        ["products",  "Produktkatalog mit heterogenen Spezifikationen"],
        ["baskets",   "Warenkörbe (Referenzen auf Produkte)"],
        ["wishlists", "Wunschlisten"],
        ["orders",    "Bestellhistorie mit Snapshot-Daten (Preis zum Kaufzeitpunkt)"],
        ["reviews",   "Produktbewertungen (verified purchase)"],
    ],
    col_widths=[4, 12.5]
)

heading("Backup-Befehle", level=2)

heading("Manuell (Linux / Git Bash)", level=3)

code_block([
    "# Backup mit Zeitstempel erstellen (readonly-User = Least Privilege)",
    "docker exec electroswap_mongo mongodump \\",
    "  --username electroswap_readonly \\",
    '  --password "ReadPass789!" \\',
    "  --authenticationDatabase electroswap \\",
    "  --db electroswap \\",
    "  --archive | gzip > ./backup/dumps/electroswap_$(date +%Y%m%d_%H%M%S).gz",
])

heading("Backup-Skript (Linux / Git Bash)", level=3)

code_block([
    "# Skript ausführen – erstellt Backup mit Zeitstempel, bereinigt >7 Tage",
    "bash demo_backup.sh",
    "",
    "# Fehlschlagendes Backup (Demo – falsches Passwort)",
    "bash demo_backup_fail.sh",
])

heading("Backup-Skript (Windows PowerShell)", level=3)

code_block([
    "# Backup erstellen – schreibt in .\\backup\\dumps\\electroswap_TIMESTAMP\\",
    ".\\backup.ps1",
])

heading("Restore-Befehle", level=2)

heading("Manuell (Linux / Git Bash)", level=3)

code_block([
    "# Restore aus komprimiertem Archiv (Admin-User erforderlich für --drop)",
    "gunzip -c ./backup/dumps/electroswap_20260321_120000.gz \\",
    "  | docker exec -i electroswap_mongo mongorestore \\",
    "      --username electroswap_admin \\",
    '      --password "AdminPass123!" \\',
    "      --authenticationDatabase electroswap \\",
    "      --db electroswap \\",
    "      --drop \\",
    "      --archive",
    "",
    "# --drop: Löscht bestehende Collections vor dem Wiedereinspielen",
    "#         (verhindert doppelte Dokumente)",
])

heading("Restore-Skript (Linux / Git Bash)", level=3)

code_block([
    "# Restore – nimmt neuestes Backup automatisch",
    "bash demo_restore.sh",
    "",
    "# Restore – bestimmtes Archiv angeben",
    "bash demo_restore.sh backup/dumps/electroswap_20260321_120000.gz",
])

heading("Restore-Skript (Windows PowerShell)", level=3)

code_block([
    "# Restore – Ordnernamen als Parameter",
    ".\\restore.ps1 -BackupName electroswap_20260321_120000",
])

heading("Warum Readonly für Backup – Least-Privilege-Prinzip", level=2)

para(
    "Das Least-Privilege-Prinzip besagt: Jeder Prozess/Benutzer erhält nur die "
    "minimal nötigen Rechte. Ein Backup-Prozess liest die Daten nur – er muss "
    "nichts schreiben oder löschen. Daher wird electroswap_readonly verwendet. "
    "Im Fehlerfall (z.B. kompromittiertes Backup-Skript) kann kein Schaden "
    "an der Datenbank entstehen."
)

code_block([
    "# Readonly-User kann KEINEN Restore durchführen (--drop schlägt fehl)",
    "docker exec electroswap_mongo mongorestore \\",
    "  --username electroswap_readonly \\",
    '  --password "ReadPass789!" \\',
    "  --authenticationDatabase electroswap \\",
    "  --db electroswap --drop /dump/test/electroswap",
    "# → MongoServerError: not authorized on electroswap to execute command { drop: ... }",
])

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 3.4 – LIVE-DEMO BACKUP & RESTORE
# ═════════════════════════════════════════════════════════════════════════════

live_demo_banner("3.4", "Backup und Restore zeigen – mit Authentifizierung")

warning_box(
    "BILDSCHIRMAUFNAHME STARTEN! Zeige Terminal und Browser nebeneinander wenn möglich. "
    "Lies jeden Befehl laut vor oder erkläre was du tust."
)

info_box(
    "Ziel: (1) Backup mit falschem Passwort → Fehler, (2) Backup mit richtigem "
    "Passwort → Erfolg, (3) Daten löschen, (4) Restore → Daten zurück, "
    "(5) Zeigen dass Readonly nicht restoren kann."
)

heading("Vorbereitung (vor Aufnahme)", level=2)

code_block([
    "# Stack muss laufen",
    "docker compose up -d",
    "docker compose ps",
    "",
    "# Seed-Daten sicherstellen",
    "docker exec electroswap_app1 python seed_data.py",
    "",
    "# Terminal im Projektverzeichnis öffnen",
    "cd C:\\Projects\\ElectroSwap\\ElectroSwap",
])

heading("Demo-Ablauf", level=2)

demo_step(1, "Aktuellen Datenbankstand anzeigen",
    note="Zeige: Datenbank ist gefüllt – das ist der Ausgangszustand",
    commands=[
        "docker exec electroswap_mongo mongosh electroswap \\",
        '  --username electroswap_admin --password "AdminPass123!" \\',
        "  --authenticationDatabase electroswap \\",
        "  --eval 'print(\"products: \" + db.products.countDocuments() +",
        '          \" | users: \" + db.users.countDocuments() +',
        '          \" | orders: \" + db.orders.countDocuments())',
        "",
        "# Erwartete Ausgabe: products: 35 | users: 36 | orders: 35",
    ]
)

demo_step(2, "Backup mit FALSCHEM Passwort – Auth schlägt fehl",
    note="Zeige: MongoDB verweigert Zugriff → Authentifizierung schützt die Daten",
    commands=[
        "bash demo_backup_fail.sh",
        "",
        "# Erwartete Ausgabe (getestet – funktioniert):",
        "#   Failed: can't create session: ... auth error: unable to authenticate",
        "#   using mechanism 'SCRAM-SHA-256': (AuthenticationFailed) Authentication failed.",
        "#   ✓ MongoDB verweigert Zugriff",
    ]
)

demo_step(3, "Backup mit KORREKTEM Passwort – readonly-User",
    note="Zeige: Backup erstellt, Datei sichtbar in backup/dumps/",
    commands=[
        "bash demo_backup.sh",
        "",
        "# Erwartete Ausgabe (getestet – funktioniert):",
        "#   Zustand VOR: products: 35 | users: 36 | orders: 35",
        "#   done dumping electroswap.products (35 documents)",
        "#   done dumping electroswap.users (36 documents) ...",
        "#   ✓ Backup erfolgreich: ./backup/dumps/electroswap_TIMESTAMP.gz  ~12K",
        "",
        "# Dateigrösse anzeigen:",
        "ls -lh backup/dumps/",
    ]
)

demo_step(4, "Datenverlust simulieren – alle Produkte löschen",
    note="Zeige: Datenbank kaputt (0 Produkte), Browser-Katalog ist leer",
    commands=[
        "docker exec electroswap_mongo mongosh electroswap \\",
        '  --username electroswap_admin --password "AdminPass123!" \\',
        "  --authenticationDatabase electroswap \\",
        "  --eval 'db.products.deleteMany({}); print(db.products.countDocuments())'",
        "# → 0   (Datenverlust!)",
        "",
        "# Browser: http://localhost/products/ → Katalog ist jetzt leer",
    ]
)

demo_step(5, "Restore durchführen – Daten kommen zurück",
    note="Zeige: Skript zeigt Vorher/Nachher automatisch, Admin-User für --drop",
    commands=[
        "bash demo_restore.sh",
        "",
        "# Erwartete Ausgabe (getestet – funktioniert):",
        "#   Benutzer: electroswap_admin (dbOwner – für --drop nötig)",
        "#   ┌─ Zustand VOR:  products: 0 | users: 36 | ...",
        "#   finished restoring electroswap.products (35 documents, 0 failures)",
        "#   ┌─ Zustand NACH: products: 35 | users: 36 | ...",
        "#   ✓ Restore abgeschlossen",
    ]
)

demo_step(6, "Wiederherstellung im Browser zeigen",
    note="Zeige: Produktkatalog ist vollständig zurück",
    commands=[
        "# Browser: http://localhost/products/",
        "# → Alle 35 Produkte sind wieder sichtbar",
    ]
)

demo_step(7, "Readonly-User kann nicht restoren (Least Privilege)",
    note="Zeige: Berechtigungstrennung – nur Admin darf --drop",
    commands=[
        "docker exec electroswap_mongo mongorestore \\",
        "  --username electroswap_readonly \\",
        '  --password "ReadPass789!" \\',
        "  --authenticationDatabase electroswap \\",
        "  --db electroswap --drop \\",
        "  /dump/electroswap_20260321_120000/electroswap",
        "# → MongoServerError: not authorized on electroswap to execute command { drop: ... }",
    ]
)

warning_box("Bildschirmaufnahme STOPPEN. Alle 7 Schritte + 3 Skripte erfolgreich gezeigt.")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 3.5 – KONZEPT HORIZONTALE SKALIERUNG (DOKU)
# ═════════════════════════════════════════════════════════════════════════════

heading("3.5  Konzept für horizontale Skalierung", level=1)

info_box("Aufgabe: Konzept für horizontale Skalierung dokumentieren.")

heading("Problemstellung", level=2)

para(
    "Eine einzelne Flask-Instanz bildet einen Single Point of Failure (SPOF): "
    "Fällt sie aus, ist die gesamte Applikation nicht erreichbar. "
    "Ausserdem ist sie auf die Ressourcen eines einzigen Servers begrenzt. "
    "Horizontale Skalierung löst beide Probleme durch mehrere parallele App-Instanzen."
)

heading("Architekturübersicht", level=2)

code_block([
    "                       ┌─────────────────┐",
    "         Internet  →   │     Nginx       │  (Load Balancer, Port 80)",
    "                       │  Round-Robin    │",
    "                       └────────┬────────┘",
    "                  ┌─────────────┼─────────────┐",
    "                  ▼             ▼             ▼",
    "            ┌──────────┐  ┌──────────┐  ┌──────────┐",
    "            │  app1    │  │  app2    │  │  app3    │",
    "            │  :5000   │  │  :5000   │  │  :5000   │",
    "            └────┬─────┘  └────┬─────┘  └────┬─────┘",
    "                 └─────────────┼─────────────┘",
    "                               ▼",
    "                       ┌───────────────┐",
    "                       │   MongoDB     │",
    "                       │   :27017      │",
    "                       └───────────────┘",
])

heading("Warum ist horizontale Skalierung möglich?", level=2)

para("ElectroSwap ist zustandslos (stateless) auf Applikationsebene. "
     "Das ist die Voraussetzung für horizontale Skalierung:")

simple_table(
    ["Aspekt", "Lösung", "Begründung"],
    [
        ["Sessions",       "Signierte Client-seitige Cookies (Flask SECRET_KEY)",
                           "Keine serverseitige Session-Speicherung nötig"],
        ["Persistenz",     "Alle App-Nodes verbinden sich mit derselben MongoDB",
                           "Kein lokaler Zustand – jeder Node kann jeden Request verarbeiten"],
        ["Dateisystem",    "Keine lokalen Datei-Uploads",
                           "Produktbilder via URL – kein geteilter Dateispeicher nötig"],
        ["Cache",          "Kein app-lokaler Cache",
                           "Keine Inkonsistenz zwischen den Nodes"],
    ],
    col_widths=[3.5, 6, 7]
)

heading("Komponenten", level=2)

simple_table(
    ["Komponente", "Technologie", "Funktion"],
    [
        ["Load Balancer", "Nginx 1.25",      "Verteilt HTTP-Requests Round-Robin auf 3 App-Nodes"],
        ["App-Nodes",     "Flask (3×) Docker", "Verarbeiten Requests unabhängig voneinander"],
        ["Datenbank",     "MongoDB 7.0",     "Gemeinsamer, persistenter Datenspeicher"],
    ],
    col_widths=[4, 4.5, 8]
)

heading("Nginx Load Balancing – Round-Robin", level=2)

para("Konfiguration in nginx/nginx.conf:")

code_block([
    "upstream electroswap_app {",
    "    server app1:5000;",
    "    server app2:5000;",
    "    server app3:5000;",
    "}",
    "",
    "server {",
    "    listen 80;",
    "    location / {",
    "        proxy_pass         http://electroswap_app;",
    "        proxy_set_header   Host            $host;",
    "        proxy_set_header   X-Real-IP       $remote_addr;",
    "        proxy_set_header   X-Forwarded-For $proxy_add_x_forwarded_for;",
    "    }",
    "}",
])

para(
    "Round-Robin: Jeder neue HTTP-Request wird reihum an den nächsten Node "
    "weitergeleitet (Request 1 → app1, Request 2 → app2, Request 3 → app3, "
    "Request 4 → app1, ...)."
)

heading("Skalierbarkeit", level=2)

simple_table(
    ["Szenario", "Massnahme"],
    [
        ["Mehr Last",      "Weiteren App-Service in docker-compose.yml + Nginx-Upstream ergänzen"],
        ["Node fällt aus", "Nginx erkennt dies und leitet automatisch nur noch an gesunde Nodes weiter"],
        ["DB-Engpass",     "MongoDB Replica Set oder Sharding einrichten (nächste Stufe)"],
    ],
    col_widths=[5, 11.5]
)

heading("Replikation und Partitionierung", level=2)

heading("Replikation (Replica Set)", level=3)
para(
    "MongoDB Replica Sets spiegeln Daten auf mehrere Nodes. "
    "Ein Primary-Node nimmt Schreiboperationen entgegen, Secondary-Nodes "
    "replizieren automatisch. Fällt der Primary aus, wählen die Secondaries "
    "automatisch einen neuen Primary (Failover)."
)

heading("Sharding (Partitionierung)", level=3)
para(
    "Bei sehr grossen Datenmengen wird die Datenbank horizontal partitioniert. "
    "Die Daten werden anhand eines Shard-Keys auf mehrere Shard-Nodes aufgeteilt. "
    "Ein Config-Server und Mongos-Router koordinieren die Anfragen."
)

heading("Einschränkungen der Demo-Architektur", level=2)

simple_table(
    ["Einschränkung", "Beschreibung"],
    [
        ["MongoDB Single Node",  "Kein Replica Set → Datenbankebene ist nicht redundant"],
        ["Gleicher Host",        "Alle Nodes laufen auf demselben Rechner → kein echter Hardware-Ausfall"],
        ["Kein HTTPS",           "TLS-Termination bei Nginx wäre für Produktion notwendig"],
    ],
    col_widths=[5, 11.5]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 3.6 – LIVE-DEMO HORIZONTALE SKALIERUNG
# ═════════════════════════════════════════════════════════════════════════════

live_demo_banner("3.6", "Horizontale Skalierung mit 3 Nodes realisieren")

warning_box(
    "BILDSCHIRMAUFNAHME STARTEN! Terminal und Browser sichtbar halten. "
    "Bei jedem Schritt die Ausgabe kurz kommentieren."
)

info_box(
    "Ziel: (1) Stack mit 3 Nodes starten, (2) Load Balancing verifizieren, "
    "(3) Node-Ausfall simulieren, (4) Authentifizierung der Nodes zeigen, "
    "(5) Seed-Daten und App im Browser zeigen."
)

heading("Vorbereitung (vor Aufnahme)", level=2)

code_block([
    "# Sicherstellen dass kein alter Stack läuft",
    "docker compose down",
    "",
    "# Projektverzeichnis",
    "cd C:\\Projects\\ElectroSwap\\ElectroSwap",
    "",
    "# Docker Desktop muss laufen",
])

heading("Demo-Ablauf", level=2)

demo_step(1, "Stack starten – alle 5 Services hochfahren",
    note="Zeige: docker-compose.yml – 3 App-Nodes, 1 Nginx, 1 MongoDB",
    commands=[
        "# Stack bauen und starten (zeige Terminal-Output)",
        "docker compose up --build",
        "",
        "# Abwarten bis alle Services 'Started' / 'healthy' sind",
        "# Erkläre dabei: app1, app2, app3 sind identische Flask-Instanzen",
    ]
)

demo_step(2, "Laufende Services anzeigen",
    note="Zeige: 5 Container laufen – Nginx, app1, app2, app3, MongoDB",
    commands=[
        "# Neues Terminal:",
        "docker compose ps",
        "",
        "# Erwartete Ausgabe:",
        "# NAME                  STATUS          PORTS",
        "# electroswap_nginx     Up              0.0.0.0:80->80/tcp",
        "# electroswap_app1      Up              5000/tcp",
        "# electroswap_app2      Up              5000/tcp",
        "# electroswap_app3      Up              5000/tcp",
        "# electroswap_mongo     Up (healthy)    0.0.0.0:27017->27017/tcp",
    ]
)

demo_step(3, "Seed-Daten laden und App im Browser zeigen",
    note="Zeige: App läuft, Produkte sind sichtbar",
    commands=[
        "docker exec electroswap_app1 python seed_data.py",
        "",
        "# Browser: http://localhost",
        "# → Startseite mit Produkten zeigen",
        "# → Login als admin@electroswap.ch / admin123",
        "# → Admin-Panel zeigen",
    ]
)

demo_step(4, "Load Balancing verifizieren – Requests auf 3 Nodes verteilt",
    note="Zeige: Requests gehen reihum an app1, app2, app3",
    commands=[
        "# Terminal 1 – Logs beobachten:",
        "docker compose logs -f app1 app2 app3",
        "",
        "# Terminal 2 – 9 Requests senden:",
        "for i in $(seq 1 9); do curl -s -o /dev/null -w \"%{http_code}\\n\" http://localhost/; done",
        "",
        "# In den Logs erkennbar:",
        "#   app1 | GET / HTTP/1.0",
        "#   app2 | GET / HTTP/1.0",
        "#   app3 | GET / HTTP/1.0",
        "#   → Round-Robin Verteilung",
    ]
)

demo_step(5, "Node-Ausfall simulieren",
    note="Zeige: App bleibt erreichbar wenn ein Node ausfällt",
    commands=[
        "# app2 stoppen (Ausfall simulieren)",
        "docker compose stop app2",
        "",
        "docker compose ps",
        "# → app2 ist 'Exited'",
        "",
        "# App ist noch erreichbar über app1 und app3:",
        "curl http://localhost/",
        "# → 200 OK",
        "",
        "# Logs zeigen: nur noch app1 und app3 empfangen Requests",
        "docker compose logs -f app1 app3",
        "",
        "# app2 wieder starten:",
        "docker compose start app2",
        "docker compose ps",
        "# → app2 ist wieder 'Up'",
    ]
)

demo_step(6, "Datenbank-Authentifizierung der App-Nodes zeigen",
    note="Zeige: Alle Nodes verbinden sich mit Auth als electroswap_app",
    commands=[
        "# Verbindungsstring des App-Nodes anzeigen:",
        "docker exec electroswap_app1 env | grep MONGO_URI",
        "# → mongodb://electroswap_app:AppPass456!@mongo:27017/electroswap?authSource=electroswap",
        "",
        "# Verbindung direkt testen (readWrite-User):",
        "docker exec -it electroswap_mongo mongosh \\",
        '  "mongodb://electroswap_app:AppPass456!@localhost/electroswap?authSource=electroswap" \\',
        '  --eval "db.products.countDocuments()"',
        "# → Anzahl Produkte (z.B. 35)",
    ]
)

demo_step(7, "Stack stoppen",
    note="Zeige: Sauberes Beenden",
    commands=[
        "# Stoppen (Daten bleiben erhalten)",
        "docker compose down",
        "",
        "# Oder: Kompletter Reset inkl. Volumes",
        "# docker compose down -v",
    ]
)

warning_box("Bildschirmaufnahme STOPPEN. Alle 7 Schritte abgeschlossen.")

# ═════════════════════════════════════════════════════════════════════════════
# SPEICHERN
# ═════════════════════════════════════════════════════════════════════════════

output_path = "docs/Kapitel3_Dokumentation_v2.docx"
doc.save(output_path)
print(f"Dokument gespeichert: {output_path}")
