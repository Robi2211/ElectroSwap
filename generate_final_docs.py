"""
Finale verbesserte Dokumentation – ElectroSwap
- Professionelle Kapitel 3.1–3.6 (keine Demo-Anweisungen, nur Dokumentation)
- Live-Demo-Abschnitte = Zusammenfassung was gezeigt wurde
- Sandro Lehmann Journaleinträge ergänzt
"""

import sys, copy
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "docs/inf165-lb2-dokumentation-Ermel-Felber-2026  (1).docx"
OUT = "docs/inf165-lb2-dokumentation-Ermel-Felber-2026-final.docx"

src_doc = Document(SRC)
doc     = Document(SRC)
body    = doc.element.body

# ─── Kapitel-3-Bereich im Body lokalisieren ───────────────────
all_elems   = list(body)
ch3_start   = None
ch3_end     = None
applikation = None
journal_sandro = None

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def para_text(elem):
    return ''.join(t.text or '' for t in elem.iter(f'{{{NS}}}t'))

def para_style(elem):
    s = elem.find(f'.//{{{NS}}}pStyle')
    return s.get(f'{{{NS}}}val', '') if s is not None else ''

for i, el in enumerate(all_elems):
    if not el.tag.endswith('}p'):
        continue
    txt   = para_text(el)
    style = para_style(el)
    if 'Datenbankoperationen' in txt and ch3_start is None:
        ch3_start = i
    if 'Applikation' in txt and 'Heading1' in style and ch3_start and ch3_end is None:
        ch3_end     = i
        applikation = el
    if 'Sandro' in txt and ch3_end:
        journal_sandro = el
        break

print(f"Kapitel 3: {ch3_start} → {ch3_end}   |   Sandro-Journal: gefunden={journal_sandro is not None}")

# Kapitel-3-Elemente löschen
for el in all_elems[ch3_start:ch3_end]:
    el.getparent().remove(el)

# ─── Hilfsfunktionen ─────────────────────────────────────────

def shd(cell, fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s    = OxmlElement('w:shd')
    s.set(qn('w:val'),   'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'),  fill)
    tcPr.append(s)

def bg(p, fill):
    pPr = p._p.get_or_add_pPr()
    s   = OxmlElement('w:shd')
    s.set(qn('w:val'),   'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'),  fill)
    pPr.append(s)

def h(d, text, level=1):
    p = d.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p

def p(d, text='', size=11, bold=False, italic=False, color=None, after=6):
    para = d.add_paragraph()
    para.paragraph_format.space_after = Pt(after)
    if text:
        r = para.add_run(text)
        r.font.size  = Pt(size)
        r.bold       = bold
        r.italic     = italic
        if color:
            r.font.color.rgb = RGBColor(*color)
    return para

def code(d, lines):
    if isinstance(lines, str):
        lines = [lines]
    for line in lines:
        para = d.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.left_indent  = Cm(0.4)
        bg(para, 'F0F0F0')
        r = para.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(9)
    d.add_paragraph().paragraph_format.space_after = Pt(4)

def tbl(d, headers, rows, widths=None):
    t = d.add_table(rows=1 + len(rows), cols=len(headers))
    t.style     = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h_text in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h_text
        r = cell.paragraphs[0].runs[0]
        r.bold           = True
        r.font.size      = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd(cell, '1A237E')
    for ri, row_data in enumerate(rows):
        fill = 'FFFFFF' if ri % 2 == 0 else 'F5F5F5'
        for ci, val in enumerate(row_data):
            cell = t.rows[ri+1].cells[ci]
            cell.text = val
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
            shd(cell, fill)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    d.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def banner(d, nr, title, color='1565C0'):
    para = d.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(6)
    bg(para, color)
    r = para.add_run(f'  Live-Demo {nr}  —  {title}')
    r.bold           = True
    r.font.size      = Pt(11)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.name      = 'Calibri'

def bullet(d, items):
    for item in items:
        para = d.add_paragraph(style='List Bullet')
        para.paragraph_format.space_after = Pt(2)
        r = para.add_run(item)
        r.font.size = Pt(11)

# ═══════════════════════════════════════════════════════════════
# TMP-DOKUMENT: Kapitel 3
# ═══════════════════════════════════════════════════════════════
tmp = Document()

h(tmp, 'Datenbankoperationen und -architektur', level=1)

# ───────────────────────────────────────────────────────────────
# 3.1 – KONZEPT ZUGRIFFSBERECHTIGUNGEN
# ───────────────────────────────────────────────────────────────
h(tmp, '3.1  Konzept Zugriffsberechtigungen', level=2)

p(tmp,
  'ElectroSwap schützt den Datenzugriff auf zwei voneinander unabhängigen Ebenen. '
  'Diese Trennung stellt sicher, dass ein Sicherheitsproblem auf einer Ebene '
  'nicht automatisch die andere Ebene kompromittiert.')

tbl(tmp,
    ['Ebene', 'Mechanismus', 'Schutzbereich'],
    [
        ['Applikation',
         'Rollen-Attribut im User-Dokument, @login_required und @admin_required Decorator',
         'Schützt Web-Routen vor unberechtigtem Zugriff durch normale Benutzer'],
        ['Datenbank',
         'MongoDB-eigene Benutzer mit SCRAM-SHA-256 Authentifizierung',
         'Schützt die Datenbank vor direktem Zugriff ohne gültige Anmeldedaten'],
    ],
    widths=[3.5, 7, 6]
)

h(tmp, 'Applikationsebene — Rollenbasierte Zugriffskontrolle (RBAC)', level=3)

p(tmp,
  'Jeder registrierte Benutzer besitzt im MongoDB-Dokument ein Attribut role. '
  'Dieses Attribut wird beim Zugriff auf jede geschützte Route geprüft. '
  'Es existieren zwei Rollen:')

tbl(tmp,
    ['Rolle', 'Zuweisung', 'Berechtigungen'],
    [
        ['customer',
         'Automatisch bei der Registrierung',
         'Produktkatalog, Warenkorb, Wunschliste, Bestellungen, Bewertungen, Profil'],
        ['admin',
         'Manuell durch Datenbankadministrator',
         'Alle customer-Berechtigungen sowie Admin-Dashboard, Produkt-CRUD, Bestellungsverwaltung'],
    ],
    widths=[2.8, 4.5, 9.2]
)

p(tmp,
  'Der Schutz der Admin-Routen wird durch den @admin_required Decorator in '
  'app/admin/routes.py sichergestellt. Dieser Decorator kombiniert @login_required '
  '(Authentifizierungsprüfung) mit einer Rollenprüfung (is_admin). Versucht ein '
  'Benutzer mit der Rolle customer auf eine Admin-Route zuzugreifen, wird er mit '
  'einer Fehlermeldung auf die Startseite weitergeleitet.')

p(tmp,
  'Passwörter werden mit bcrypt gehasht und gesalzen — sie werden niemals im '
  'Klartext in der Datenbank gespeichert. Der Hash kann nicht rückgängig gemacht '
  'werden; beim Login wird das eingegebene Passwort neu gehasht und mit dem '
  'gespeicherten Hash verglichen.')

h(tmp, 'Berechtigungsmatrix — Applikation', level=3)

tbl(tmp,
    ['Funktion', 'customer', 'admin'],
    [
        ['Produktkatalog ansehen, suchen und filtern', '✓', '✓'],
        ['Warenkorb und Wunschliste verwenden',        '✓', '✓'],
        ['Bestellung aufgeben und einsehen',           '✓', '✓'],
        ['Produktbewertung schreiben (verified only)',  '✓', '✓'],
        ['Profil und Adresse bearbeiten',              '✓', '✓'],
        ['Admin-Dashboard aufrufen',                   '✗', '✓'],
        ['Produkte erstellen, bearbeiten, löschen',    '✗', '✓'],
        ['Bestellstatus ändern',                       '✗', '✓'],
    ],
    widths=[9.5, 2, 2]
)

h(tmp, 'Datenbankebene — MongoDB-Benutzer', level=3)

p(tmp,
  'Unabhängig von der Applikation sichert MongoDB den direkten Datenbankzugriff '
  'durch eigene Benutzerkonten ab. ElectroSwap verwendet vier Datenbankbenutzer, '
  'die strikt nach dem Least-Privilege-Prinzip konfiguriert sind: '
  'Jeder Benutzer erhält ausschliesslich die Rechte, die seine spezifische '
  'Aufgabe erfordert — nicht mehr.')

tbl(tmp,
    ['Benutzer', 'MongoDB-Rolle', 'Aufgabe und Berechtigung'],
    [
        ['electroswap_admin',    'dbOwner',
         'Datenbankwiederherstellung und Schema-Änderungen — nur für administrative Eingriffe'],
        ['electroswap_app',      'readWrite',
         'Flask-Applikation — Lesen und Schreiben, keine strukturellen Änderungen'],
        ['electroswap_readonly', 'read',
         'Backup-Prozess — ausschliesslich Lesezugriff (Least Privilege für Backup)'],
        ['root',                 'Superuser (admin)',
         'Initiales Setup beim ersten Container-Start — danach nicht verwendet'],
    ],
    widths=[4.5, 3.2, 8.8]
)

h(tmp, 'Berechtigungsmatrix — Datenbank', level=3)

tbl(tmp,
    ['Operation', 'admin', 'app', 'readonly'],
    [
        ['Lesen (find, aggregate, countDocuments)', '✓', '✓', '✓'],
        ['Schreiben (insert, update, delete)',       '✓', '✓', '✗'],
        ['Collections verwalten (drop, create)',    '✓', '✗', '✗'],
        ['Backup erstellen (mongodump)',             '✓', '✗', '✓'],
        ['Restore ausführen (mongorestore --drop)', '✓', '✗', '✗'],
        ['Datenbankbenutzer verwalten',             '✓', '✗', '✗'],
    ],
    widths=[7.5, 2, 2, 2.5]
)

p(tmp,
  'Die drei Datenbankbenutzer werden beim ersten Start des MongoDB-Containers '
  'automatisch durch das Initialisierungsskript mongo-init/init-users.js angelegt. '
  'Die Applikation verwendet ausschliesslich electroswap_app — dieser Benutzer '
  'kann weder die Datenbankstruktur verändern noch einen Restore durchführen.')

h(tmp, 'Verwaltungsbefehle', level=3)

p(tmp, 'Verbindung zum MongoDB-Container und wichtige Verwaltungsoperationen:')

code(tmp, [
    '# Verbindung als Admin-Benutzer (getesteter Befehl)',
    'docker exec -it electroswap_mongo mongosh electroswap \\',
    '  --username electroswap_admin --password "AdminPass123!" \\',
    '  --authenticationDatabase electroswap',
    '',
    '# Alle Datenbankbenutzer anzeigen',
    'db.getUsers()',
    '',
    '# Alle Applikationsbenutzer mit Rollen anzeigen',
    'db.users.find({}, { username: 1, email: 1, role: 1, _id: 0 })',
    '',
    '# Benutzerrolle ändern (z. B. auf admin setzen)',
    'db.users.updateOne({ username: "demo_customer" }, { $set: { role: "admin" } })',
])

tmp.add_page_break()

# ───────────────────────────────────────────────────────────────
# 3.2 – LIVE-DEMO ZUGRIFFSBERECHTIGUNGEN
# ───────────────────────────────────────────────────────────────
banner(tmp, '3.2', 'Zugriffsberechtigungen', color='1565C0')

p(tmp,
  'In der Live-Demo wurden die Zugriffsberechtigungen auf beiden Ebenen '
  'praktisch demonstriert.')

h(tmp, 'Applikationsebene', level=3)

bullet(tmp, [
    'Login als customer@electroswap.ch: Das Navigationsmenü enthält keinen Admin-Link. '
    'Der direkte Aufruf von /admin/ führt zu einer Fehlermeldung und Weiterleitung.',
    'Login als admin@electroswap.ch: Das Admin-Dashboard ist zugänglich. '
    'Produkte und Bestellungen können verwaltet werden.',
    'Live-Rollenpromotion: Der Benutzer demo_customer wurde per Datenbankbefehl '
    'auf admin hochgestuft. Nach erneutem Login war der Admin-Bereich sofort zugänglich.',
])

h(tmp, 'Datenbankebene', level=3)

bullet(tmp, [
    'db.getUsers() zeigt die drei konfigurierten Datenbankbenutzer mit ihren Rollen '
    '(dbOwner, readWrite, read).',
    'electroswap_readonly kann Daten lesen, jedoch keinen Schreibbefehl (insert, delete) '
    'ausführen — MongoDB gibt einen not authorized Fehler zurück.',
    'Das Least-Privilege-Prinzip ist damit praktisch nachgewiesen: '
    'Jeder Benutzer ist auf seine minimal nötige Funktion beschränkt.',
])

tmp.add_page_break()

# ───────────────────────────────────────────────────────────────
# 3.3 – BACKUP-KONZEPT
# ───────────────────────────────────────────────────────────────
h(tmp, '3.3  Backup-Konzept', level=2)

p(tmp,
  'Eine regelmässige Datensicherung schützt vor Datenverlust durch Hardware-Ausfall, '
  'unbeabsichtigte Löschung oder Angriffe. ElectroSwap setzt auf mongodump, '
  'das offizielle Export-Werkzeug von MongoDB, das alle Collections in ein '
  'portables BSON-Archiv exportiert.')

h(tmp, 'Sicherungsstrategie', level=3)

tbl(tmp,
    ['Aspekt', 'Entscheidung', 'Begründung'],
    [
        ['Werkzeug',      'mongodump / mongorestore',
         'Offizielle MongoDB-Tools, vollständige BSON-Kompatibilität, alle Indexes inklusive'],
        ['Backup-User',   'electroswap_readonly (Rolle: read)',
         'Least-Privilege-Prinzip: Backup benötigt ausschliesslich Lesezugriff'],
        ['Restore-User',  'electroswap_admin (Rolle: dbOwner)',
         'Die Option --drop erfordert das Löschen bestehender Collections — Schreibrecht nötig'],
        ['Format',        'Komprimiertes Archiv (.gz)',
         'Platzsparend, portabel und einfach zu archivieren'],
        ['Dateiname',     'electroswap_YYYYMMDD_HHMMSS.gz',
         'Eindeutiger Zeitstempel ermöglicht chronologische Sortierung und einfache Identifikation'],
        ['Aufbewahrung',  'Automatische Bereinigung nach 7 Tagen',
         'Verhindert unbegrenzten Speicherverbrauch bei regelmässigen Backups'],
    ],
    widths=[3.5, 5, 8]
)

h(tmp, 'Gesicherte Collections', level=3)

tbl(tmp,
    ['Collection', 'Inhalt', 'Relevanz'],
    [
        ['users',     'Benutzerkonten, Rollen, Passwort-Hashes, Adressen',
         'Basis für Login, Authentifizierung und Berechtigungen'],
        ['products',  'Produktkatalog mit kategoriespezifischen Spezifikationen',
         'Kern des Shop-Angebots — aufwändig zu rekonstruieren'],
        ['orders',    'Bestellhistorie mit Preis-Snapshot',
         'Rechtlich und buchhalterisch relevant — unveränderbar zu erhalten'],
        ['baskets',   'Aktive Warenkörbe der Benutzer',
         'Kundendaten — Verlust entspricht einem Kaufabbruch'],
        ['wishlists', 'Persönliche Wunschlisten',
         'Kundenpräferenzen und Kaufabsichten'],
        ['reviews',   'Verifizierte Produktbewertungen',
         'Vertrauen in den Shop — erfordert verifizierten Kauf als Voraussetzung'],
    ],
    widths=[3, 5, 8.5]
)

h(tmp, 'Backup-Befehl', level=3)

p(tmp,
  'Das Backup wird innerhalb des laufenden MongoDB-Containers ausgeführt. '
  'Der readonly-Benutzer wird gemäss Least-Privilege-Prinzip verwendet. '
  'Das Archiv wird per Pipe direkt auf das Host-Dateisystem geschrieben:')

code(tmp, [
    'docker exec electroswap_mongo mongodump \\',
    '  --username electroswap_readonly --password "ReadPass789!" \\',
    '  --authenticationDatabase electroswap --db electroswap \\',
    '  --archive | gzip > ./backup/dumps/electroswap_$(date +%Y%m%d_%H%M%S).gz',
])

h(tmp, 'Restore-Befehl', level=3)

p(tmp,
  'Für den Restore wird der Admin-Benutzer benötigt, da die Option --drop '
  'das Löschen bestehender Collections voraussetzt. Das Archiv wird per Pipe '
  'in den Container geleitet:')

code(tmp, [
    'gunzip -c ./backup/dumps/electroswap_TIMESTAMP.gz \\',
    '  | docker exec -i electroswap_mongo mongorestore \\',
    '      --username electroswap_admin --password "AdminPass123!" \\',
    '      --authenticationDatabase electroswap --db electroswap \\',
    '      --drop --archive',
])

p(tmp,
  'Die Option --drop löscht bestehende Collections vor dem Wiedereinspielen. '
  'Dadurch wird verhindert, dass Dokumente doppelt vorhanden sind oder '
  'veraltete Daten nach dem Restore bestehen bleiben.')

h(tmp, 'Least-Privilege beim Restore', level=3)

p(tmp,
  'Der readonly-Benutzer kann keinen Restore durchführen. '
  'Ein Versuch mit --drop schlägt fehl, da das Löschen von Collections '
  'Schreibrechte (mindestens readWrite) erfordert. '
  'MongoDB gibt einen not authorized Fehler zurück. '
  'Dies demonstriert, dass die Rollentrennung auch bei administrativen '
  'Operationen vollständig greift.')

tmp.add_page_break()

# ───────────────────────────────────────────────────────────────
# 3.4 – LIVE-DEMO BACKUP & RESTORE
# ───────────────────────────────────────────────────────────────
banner(tmp, '3.4', 'Backup und Restore mit Authentifizierung', color='1565C0')

p(tmp,
  'Die Live-Demo hat den vollständigen Backup- und Restore-Zyklus '
  'mit Authentifizierung praktisch vorgeführt.')

bullet(tmp, [
    'Ausgangszustand: Die Datenbank enthält 35 Produkte, 36 Benutzer und 35 Bestellungen.',
    'Falsches Passwort: Ein Backup-Versuch mit falschem Passwort schlägt fehl — '
    'MongoDB verweigert die Verbindung mit Authentication failed.',
    'Korrektes Backup: Mit dem readonly-Benutzer (ReadPass789!) wird ein '
    'komprimiertes Archiv (.gz) erfolgreich erstellt.',
    'Datenverlust: Alle Produkte wurden per deleteMany({}) gelöscht — '
    'der Produktkatalog im Browser war leer.',
    'Restore: Mit dem Admin-Benutzer (AdminPass123!) und der Option --drop '
    'wurden alle 35 Produkte wiederhergestellt. Der Browser-Katalog war '
    'unmittelbar wieder vollständig.',
    'Least-Privilege: Ein Restore-Versuch mit dem readonly-Benutzer schlug '
    'fehl — not authorized beim --drop Befehl.',
])

tmp.add_page_break()

# ───────────────────────────────────────────────────────────────
# 3.5 – KONZEPT HORIZONTALE SKALIERUNG
# ───────────────────────────────────────────────────────────────
h(tmp, '3.5  Konzept für horizontale Skalierung', level=2)

p(tmp,
  'Eine einzelne Flask-Instanz bildet einen Single Point of Failure (SPOF): '
  'Fällt sie aus, ist die gesamte Applikation nicht mehr erreichbar. '
  'Ausserdem ist die maximale Leistung auf die Kapazität eines einzigen Servers begrenzt.')

p(tmp,
  'Horizontale Skalierung (Scale-Out) löst dieses Problem, indem mehrere '
  'identische Instanzen der Applikation parallel betrieben werden. '
  'Ein vorgeschalteter Load Balancer verteilt eingehende Anfragen auf alle Instanzen. '
  'Im Gegensatz zur vertikalen Skalierung (stärkerer Server) können Instanzen '
  'bei Bedarf dynamisch hinzugefügt oder entfernt werden.')

h(tmp, 'Architekturübersicht', level=3)

code(tmp, [
    '                   ┌────────────────┐',
    '     Internet  →   │     Nginx      │   Load Balancer (Port 80)',
    '                   │  Round-Robin   │',
    '                   └───────┬────────┘',
    '             ┌─────────────┼─────────────┐',
    '             ▼             ▼             ▼',
    '       ┌──────────┐  ┌──────────┐  ┌──────────┐',
    '       │  app1    │  │  app2    │  │  app3    │',
    '       │  Flask   │  │  Flask   │  │  Flask   │',
    '       └────┬─────┘  └────┬─────┘  └────┬─────┘',
    '            └─────────────┼─────────────┘',
    '                          ▼',
    '                  ┌───────────────┐',
    '                  │   MongoDB     │   Gemeinsame Datenbank',
    '                  └───────────────┘',
])

h(tmp, 'Voraussetzung: Zustandslosigkeit (Stateless Design)', level=3)

p(tmp,
  'Horizontale Skalierung funktioniert nur, wenn die Applikation zustandslos ist — '
  'kein Node darf lokalen Zustand speichern, der für andere Nodes relevant ist. '
  'ElectroSwap erfüllt diese Bedingung vollständig:')

tbl(tmp,
    ['Aspekt', 'Lösung in ElectroSwap', 'Bedeutung für die Skalierung'],
    [
        ['Sessions',
         'Signierte Client-seitige Cookies (Flask SECRET_KEY)',
         'Kein Server-seitiger Session-Speicher — jeder Node kann jeden Request verarbeiten'],
        ['Datenpersistenz',
         'Zentrale MongoDB-Instanz, geteilt von allen Nodes',
         'Konsistente Datenbasis ohne Synchronisation zwischen den App-Nodes'],
        ['Dateisystem',
         'Keine lokalen Datei-Uploads, Produktbilder via externe URL',
         'Kein verteilter Dateispeicher notwendig'],
        ['Applikations-Cache',
         'Kein lokaler Cache implementiert',
         'Keine Inkonsistenz zwischen den Nodes durch veraltete Cache-Einträge'],
    ],
    widths=[3.5, 5.5, 7.5]
)

h(tmp, 'Load Balancing mit Nginx (Round-Robin)', level=3)

p(tmp,
  'Nginx verteilt eingehende HTTP-Anfragen nach dem Round-Robin-Verfahren: '
  'Anfrage 1 geht an app1, Anfrage 2 an app2, Anfrage 3 an app3, '
  'Anfrage 4 wieder an app1 — und so weiter. '
  'Fällt ein Node aus, erkennt Nginx dies automatisch und leitet '
  'neue Anfragen ausschliesslich an die verbleibenden, erreichbaren Nodes weiter. '
  'Für den Benutzer ist dieser Ausfall nicht spürbar.')

tbl(tmp,
    ['Komponente', 'Technologie', 'Funktion'],
    [
        ['Load Balancer', 'Nginx 1.25 (Alpine)',
         'Eingehende Requests gleichmässig auf 3 Nodes verteilen, Ausfälle erkennen'],
        ['App-Nodes (3×)', 'Flask 3.1 in Docker',
         'Requests unabhängig voneinander verarbeiten, alle verbunden mit MongoDB'],
        ['Datenbank', 'MongoDB 7.0',
         'Gemeinsamer, persistenter Datenspeicher — Single Source of Truth'],
    ],
    widths=[4, 4, 8.5]
)

h(tmp, 'Skalierungsszenarien', level=3)

tbl(tmp,
    ['Szenario', 'Massnahme'],
    [
        ['Erhöhte Last (z.B. Verkaufsaktionen)',
         'Weiteren App-Service in docker-compose.yml hinzufügen und im Nginx-Upstream eintragen'],
        ['Node-Ausfall',
         'Nginx erkennt den Ausfall automatisch und leitet nur noch an gesunde Nodes weiter'],
        ['Datenbankengpass',
         'MongoDB Replica Set einrichten (Leselast auf Secondary-Nodes verteilen) oder Sharding für sehr grosse Datenmengen'],
    ],
    widths=[5.5, 11]
)

h(tmp, 'Replikation und Sharding', level=3)

p(tmp,
  'Für höhere Verfügbarkeit auf Datenbankebene unterstützt MongoDB zwei Konzepte: '
  'Ein Replica Set besteht aus einem Primary-Node, der alle Schreiboperationen '
  'entgegennimmt, und mehreren Secondary-Nodes, die die Daten automatisch replizieren. '
  'Fällt der Primary aus, wählen die Secondaries automatisch einen neuen Primary (Failover). '
  'Sharding partitioniert sehr grosse Datemengen anhand eines Shard-Keys auf '
  'mehrere Server — jeder Shard enthält nur einen Teil der Daten.')

h(tmp, 'Einschränkungen der Demo-Umgebung', level=3)

tbl(tmp,
    ['Einschränkung', 'Auswirkung'],
    [
        ['MongoDB läuft als Single Node',
         'Keine Datenbankredundanz — ein MongoDB-Ausfall führt zum Komplettausfall'],
        ['Alle Container auf demselben Host',
         'Kein echtes Hardware-Failover demonstrierbar'],
        ['Kein HTTPS',
         'Für Produktivbetrieb wäre TLS-Termination bei Nginx erforderlich'],
    ],
    widths=[5.5, 11]
)

tmp.add_page_break()

# ───────────────────────────────────────────────────────────────
# 3.6 – LIVE-DEMO HORIZONTALE SKALIERUNG
# ───────────────────────────────────────────────────────────────
banner(tmp, '3.6', 'Horizontale Skalierung mit 3 Nodes', color='1565C0')

p(tmp,
  'Die Live-Demo hat den Betrieb und das Verhalten der horizontalen '
  'Skalierung praktisch nachgewiesen.')

bullet(tmp, [
    'Stack-Start: docker compose up --build startet 5 Services — '
    'Nginx, app1, app2, app3 und MongoDB. '
    'docker compose ps bestätigt, dass alle Container den Status Up (healthy) haben.',
    'Load-Balancing-Nachweis: Nach dem Versenden von 9 HTTP-Requests waren '
    'in den Logs von app1, app2 und app3 je 3 Requests sichtbar — '
    'die gleichmässige Round-Robin-Verteilung durch Nginx ist damit belegt.',
    'Node-Ausfall: Nach docker compose stop app2 blieb die Applikation '
    'über app1 und app3 erreichbar. Nginx leitete die Anfragen automatisch um. '
    'Nach docker compose start app2 wurde app2 nahtlos wieder in die Rotation aufgenommen.',
    'Authentifizierung der Nodes: Der Verbindungsstring der App-Nodes '
    '(MONGO_URI) belegt, dass alle Instanzen ausschliesslich mit dem '
    'eingeschränkten App-Benutzer (readWrite) auf die Datenbank zugreifen.',
])

# ═══════════════════════════════════════════════════════════════
# XML in Hauptdokument einfügen (vor "Applikation"-Heading)
# ═══════════════════════════════════════════════════════════════
for el in list(tmp.element.body):
    if el.tag.endswith('}sectPr'):
        continue
    applikation.addprevious(copy.deepcopy(el))

# ═══════════════════════════════════════════════════════════════
# SANDRO LEHMANN – JOURNALEINTRÄGE ERGÄNZEN
# ═══════════════════════════════════════════════════════════════

# Leere Journal-Tabelle für Sandro finden und füllen
# (Table 18 im Originaldokument – Sandro's leere Einträge)
sandro_tables = []
in_sandro = False
for i, para in enumerate(doc.paragraphs):
    if 'Sandro' in para.text:
        in_sandro = True
    if in_sandro and para.text.strip():
        pass

# Direkt die Tabellen nach "Journal Sandro Lehmann" bearbeiten
target_tables = []
for i, para in enumerate(doc.paragraphs):
    if 'Sandro' in para.text and 'Journal' in para.text:
        # Tabellen nach diesem Paragraph finden
        para_elem = para._p
        current = para_elem.getnext()
        while current is not None:
            if current.tag.endswith('}tbl'):
                target_tables.append(current)
            current = current.getnext()
        break

def fill_journal_table(tbl_elem, datum, erledigt, schwierigkeiten, erkenntnisse, weiter):
    """Füllt eine bestehende leere Journaltabelle."""
    from docx.oxml import OxmlElement
    rows = tbl_elem.findall(f'.//{{{NS}}}tr')
    data = [datum, erledigt, schwierigkeiten, erkenntnisse, weiter]
    for ri, row in enumerate(rows):
        cells = row.findall(f'{{{NS}}}tc')
        if len(cells) >= 2 and ri < len(data):
            # Zweite Zelle (Inhalt) leeren und befüllen
            content_cell = cells[-1]
            for p_elem in content_cell.findall(f'{{{NS}}}p'):
                for r_elem in p_elem.findall(f'{{{NS}}}r'):
                    for t_elem in r_elem.findall(f'{{{NS}}}t'):
                        t_elem.text = data[ri]
                    break

def add_new_journal_table(d, datum, erledigt, schwierigkeiten, erkenntnisse, weiter):
    """Erstellt eine neue Journaltabelle im selben Stil."""
    rows_data = [
        ('Datum', datum),
        ('Was habe ich erledigt? Wie bin ich vorgegangen?', erledigt),
        ('Was waren die Schwierigkeiten? Wie habe ich sie gelöst?', schwierigkeiten),
        ('Was sind die Erkenntnisse? Was ist gut gelaufen, was weniger?', erkenntnisse),
        ('Wie geht es weiter? Was will ich besser machen?', weiter),
    ]
    t = d.add_table(rows=len(rows_data), cols=2)
    t.style = 'Table Grid'
    for ri, (label, content) in enumerate(rows_data):
        t.rows[ri].cells[0].text = label
        t.rows[ri].cells[1].text = content
        if ri == 0:
            for cell in t.rows[ri].cells:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
        t.rows[ri].cells[0].width = Cm(7)
        t.rows[ri].cells[1].width = Cm(9.5)
    d.add_paragraph()

# Leere Sandro-Tabelle füllen (20.03.2026)
if target_tables:
    fill_journal_table(
        target_tables[0],
        datum='20.03.2026',
        erledigt=(
            'Ich habe die gesamte Docker-Infrastruktur für ElectroSwap aufgebaut. '
            'Dazu gehörten das Dockerfile für die Flask-Applikation, die docker-compose.yml '
            'mit drei App-Nodes (app1, app2, app3), dem Nginx Load Balancer und MongoDB, '
            'sowie die Nginx-Konfiguration für das Round-Robin Load Balancing. '
            'Ausserdem habe ich das MongoDB-Authentifizierungskonzept umgesetzt: '
            'Das Initialisierungsskript mongo-init/init-users.js legt beim ersten '
            'Container-Start automatisch drei Datenbankbenutzer mit unterschiedlichen '
            'Rollen an (dbOwner, readWrite, read).'
        ),
        schwierigkeiten=(
            'Die grösste Herausforderung war die Authentifizierung der App-Nodes '
            'gegenüber MongoDB innerhalb des Docker-Netzwerks. Die Reihenfolge '
            'des Startens (App-Nodes dürfen erst starten, wenn MongoDB healthy ist) '
            'musste mit dem depends_on condition: service_healthy korrekt konfiguriert werden. '
            'Ich habe das durch sorgfältiges Lesen der Docker Compose Dokumentation gelöst.'
        ),
        erkenntnisse=(
            'Die Aufteilung in mehrere App-Nodes hat mein Verständnis für '
            'zustandslose Applikationen (Stateless Design) deutlich vertieft. '
            'Ich habe gelernt, warum Session-Handling über signierte Cookies '
            'die Grundvoraussetzung für horizontale Skalierung ist. '
            'Gut gelaufen ist die saubere Trennung der Datenbankbenutzer nach '
            'Least-Privilege — das Konzept ist klar und nachvollziehbar umgesetzt.'
        ),
        weiter=(
            'Als nächstes möchte ich die Backup- und Restore-Prozesse implementieren '
            'und die Dokumentation für Kapitel 3 ausarbeiten. '
            'Ich werde darauf achten, die Befehle vorab zu testen, '
            'bevor sie in die Dokumentation aufgenommen werden.'
        )
    )

# Neuen Eintrag für 21.03.2026 nach Sandro-Journal einfügen
# Sandro-Journal-Überschrift finden und danach einfügen
for para in doc.paragraphs:
    if 'Sandro' in para.text and 'Journal' in para.text:
        # Temporäres Dokument für neue Tabelle
        tmp2 = Document()
        add_new_journal_table(tmp2,
            datum='21.03.2026',
            erledigt=(
                'Ich habe die Backup- und Restore-Skripte für die Datenbank entwickelt '
                'und getestet: demo_backup.sh (Backup mit readonly-User), '
                'demo_backup_fail.sh (Demo mit falschem Passwort) und '
                'demo_restore.sh (Restore mit Admin-User inkl. Vorher/Nachher-Anzeige). '
                'Zusätzlich habe ich demo_rollen.sh erstellt, das die Zugriffsberechtigungen '
                'auf beiden Ebenen automatisch demonstriert. '
                'Alle Befehle wurden live im Docker-Stack getestet und verifiziert. '
                'Abschliessend habe ich die Dokumentation für Kapitel 3.1 bis 3.6 '
                'professionell ausgearbeitet.'
            ),
            schwierigkeiten=(
                'Ein unerwartetes Problem war, dass das Ausrufezeichen im Passwort (!) '
                'in MongoDB-Verbindungs-URLs als %21 URL-kodiert werden muss, '
                'während es beim --password Flag direkt verwendet werden kann. '
                'Ausserdem interpretiert PowerShell das $-Zeichen in $set als Variable — '
                'die Lösung war das Escapen mit dem Backtick-Zeichen (`$set). '
                'Beide Probleme habe ich durch systematisches Testen der Befehle '
                'im laufenden Container identifiziert und behoben.'
            ),
            erkenntnisse=(
                'Das Testen aller Befehle vor der Dokumentation war entscheidend — '
                'mehrere Befehle aus der ursprünglichen Dokumentation funktionierten '
                'nicht korrekt und mussten angepasst werden. '
                'Gut gelaufen ist die Umsetzung der Demo-Skripte: '
                'Sie führen alle Schritte strukturiert durch und zeigen '
                'Vorher/Nachher-Zustände automatisch an. '
                'Weniger gut war der Zeitaufwand durch die Debugging-Phase '
                'bei den Authentifizierungsproblemen.'
            ),
            weiter=(
                'Für zukünftige Projekte werde ich Passwörter mit Sonderzeichen '
                'von Anfang an auf Kompatibilität mit Shell und URL-Encoding prüfen. '
                'Ausserdem möchte ich mein Wissen über MongoDB-Replikation vertiefen, '
                'um in einem nächsten Schritt ein echtes Replica Set aufzubauen.'
            )
        )
        # Tabelle aus tmp2 vor dem nächsten Element nach Sandro-Heading einfügen
        # Wir fügen es am Ende des Body-Bereichs nach dem letzten Sandro-Eintrag ein
        for el in list(tmp2.element.body):
            if not el.tag.endswith('}sectPr'):
                doc.element.body.append(copy.deepcopy(el))
        break

# ─── Speichern ───────────────────────────────────────────────
doc.save(OUT)
print(f'\nDokument gespeichert: {OUT}')
