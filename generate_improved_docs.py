"""
Verbesserte Dokumentation – ElectroSwap Kapitel 3
Liest das Original-Dokument, kopiert alles ausser Kapitel 3,
und ersetzt Kapitel 3 (3.1–3.6) mit professionellem Inhalt.

Ausführen: python generate_improved_docs.py
"""

import sys
import copy
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Quelldokument lesen ─────────────────────────────────────
SRC = "docs/inf165-lb2-dokumentation-Ermel-Felber-2026  (1).docx"
OUT = "docs/inf165-lb2-dokumentation-Ermel-Felber-2026-verbessert.docx"

src_doc = Document(SRC)

# ─── Neues Dokument basierend auf Original ────────────────────
# Wir kopieren das Original und ersetzen nur Kapitel 3
doc = Document(SRC)

# Hilfsfunktion: Element aus Body entfernen
def remove_element(elem):
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)

# Alle Body-Elemente sammeln
body = doc.element.body
all_elements = list(body)

# ─── Kapitel 3 Bereich finden (Start + Ende) ─────────────────
chapter3_start = None
chapter3_end   = None

for i, elem in enumerate(all_elements):
    # Suche Heading 1 "Datenbankoperationen"
    if elem.tag.endswith('}p'):
        style = elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
        if style is not None and 'Heading1' in style.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val',''):
            text = ''.join(t.text or '' for t in elem.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
            if 'Datenbankoperationen' in text and chapter3_start is None:
                chapter3_start = i
            elif chapter3_start is not None and 'Applikation' in text:
                chapter3_end = i
                break

print(f"Kapitel 3: Element {chapter3_start} bis {chapter3_end}")

# Kapitel 3 Elemente entfernen
elements_to_remove = all_elements[chapter3_start:chapter3_end]
for elem in elements_to_remove:
    remove_element(elem)

# ─── Hilfsfunktionen für neues Kapitel 3 ─────────────────────

def make_shading(cell, fill_hex):
    """Hintergrundfarbe für Tabellenzelle."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_para(doc, text, size=11, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold   = bold
        r.italic = italic
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p

def add_code(doc, lines):
    """Grauer Code-Block – sparsam einsetzen."""
    if isinstance(lines, str):
        lines = [lines]
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F2F2F2")
        pPr.append(shd)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_info_box(doc, text, fill="E8F4FD"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    pPr.append(shd)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    return p

def add_demo_banner(doc, nr, title):
    """Roter Live-Demo-Banner."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "C62828")
    pPr.append(shd)
    r = p.add_run(f"  LIVE-DEMO {nr}  –  {title}  |  Bildschirmaufnahme erforderlich")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.name = "Calibri"

def add_step(doc, nr, title, desc=None):
    """Nummerierter Demo-Schritt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"Schritt {nr}  –  {title}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    if desc:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(3)
        r2 = p2.add_run(desc)
        r2.italic = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_table(doc, headers, rows, col_widths=None):
    """Formatierte Tabelle mit blauem Header."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header-Zeile
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        make_shading(cell, "1565C0")
    # Datenzeilen
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        fill = "FFFFFF" if ri % 2 == 0 else "F5F5F5"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            make_shading(cell, fill)
    # Spaltenbreiten
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

# ─── Einfügepunkt im Body finden ─────────────────────────────
# Nach dem Entfernen müssen wir die Applikation-Überschrift finden
# und das neue Kapitel 3 davor einfügen.

# Referenz-Element (Applikation-Heading) finden
ref_elem = None
for elem in list(doc.element.body):
    if elem.tag.endswith('}p'):
        text = ''.join(t.text or '' for t in elem.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
        if text.strip() == 'Applikation':
            ref_elem = elem
            break

if ref_elem is None:
    print("WARNUNG: 'Applikation' Heading nicht gefunden – füge am Ende ein")

# Hilfsfunktion: Paragraph vor ref_elem einfügen
# Wir arbeiten direkt mit dem Body und dem Referenz-Element
def insert_before_ref(new_elem):
    """Fügt new_elem direkt vor ref_elem in den Body ein."""
    if ref_elem is not None:
        ref_elem.addprevious(new_elem)
    else:
        doc.element.body.append(new_elem)

# Temporäres Dokument für Kapitel-3-Inhalt
tmp = Document()

# ════════════════════════════════════════════════════════════════
# KAPITEL 3 – DATENBANKOPERATIONEN UND ARCHITEKTUR
# ════════════════════════════════════════════════════════════════

add_heading(tmp, "Datenbankoperationen und -architektur", level=1)

# ────────────────────────────────────────────────────────────────
# 3.1 – KONZEPT ZUGRIFFSBERECHTIGUNGEN
# ────────────────────────────────────────────────────────────────
add_heading(tmp, "3.1  Zugriffsberechtigungen: Konzept", level=2)

add_para(tmp,
    "ElectroSwap schützt den Zugriff auf zwei unabhängigen Ebenen: der Applikationsebene "
    "und der Datenbankebene. Beide ergänzen sich gegenseitig – die Applikation schützt "
    "Routen vor normalen Benutzern, die Datenbank schützt die Daten vor direktem Zugriff."
)

add_table(tmp,
    ["Ebene", "Mechanismus", "Schützt vor"],
    [
        ["Applikation", "Rollen-Attribut im User-Dokument + @admin_required Decorator",
         "Unberechtigtem Zugriff auf Admin-Routen"],
        ["Datenbank",   "MongoDB-Benutzer mit SCRAM-SHA-256 Authentifizierung",
         "Direktem Datenbankzugriff ohne Berechtigung"],
    ],
    col_widths=[3.5, 7.5, 5.5]
)

add_heading(tmp, "Ebene 1: Applikationsrollen (RBAC)", level=3)

add_para(tmp,
    "Jeder Benutzer besitzt im MongoDB-Dokument ein Feld role. Dieses Feld entscheidet, "
    "welche Routen der Benutzer aufrufen darf. Es gibt zwei Rollen:"
)

add_table(tmp,
    ["Rolle", "Vergabe", "Zugang"],
    [
        ["customer", "Automatisch bei der Registrierung",
         "Produkte, Warenkorb, Bestellungen, Bewertungen, Profil"],
        ["admin",    "Manuell durch einen Datenbankadministrator",
         "Alle customer-Rechte + Admin-Dashboard, Produkt-CRUD, Bestellstatus"],
    ],
    col_widths=[3, 5, 8.5]
)

add_para(tmp,
    "Der Schutz wird im Code durch den @admin_required Decorator sichergestellt. "
    "Ruft ein customer eine Admin-Route auf, wird er mit der Meldung 'Access denied' "
    "auf die Startseite weitergeleitet. Passwörter werden niemals im Klartext gespeichert – "
    "bcrypt erzeugt einen gesalzten Hash, der nicht rückgängig gemacht werden kann."
)

add_heading(tmp, "Berechtigungsmatrix – Applikation", level=3)

add_table(tmp,
    ["Funktion", "customer", "admin"],
    [
        ["Produkte ansehen, suchen, filtern",          "✓", "✓"],
        ["Warenkorb & Wunschliste nutzen",             "✓", "✓"],
        ["Bestellungen aufgeben und einsehen",         "✓", "✓"],
        ["Produktbewertungen schreiben",               "✓", "✓"],
        ["Profil und Adresse bearbeiten",              "✓", "✓"],
        ["Admin-Dashboard aufrufen",                   "✗", "✓"],
        ["Produkte erstellen, bearbeiten, löschen",    "✗", "✓"],
        ["Bestellstatus verwalten",                    "✗", "✓"],
    ],
    col_widths=[9, 2.5, 2.5]
)

add_heading(tmp, "Ebene 2: MongoDB-Datenbankbenutzer", level=3)

add_para(tmp,
    "Unabhängig von der Applikation sichert MongoDB den direkten Datenbankzugriff ab. "
    "Es wurden vier Benutzer nach dem Least-Privilege-Prinzip angelegt: Jeder Benutzer "
    "erhält nur die Rechte, die er für seine Aufgabe wirklich benötigt."
)

add_table(tmp,
    ["Benutzer", "Rolle", "Zweck"],
    [
        ["root",                 "Superuser (admin-DB)", "Initiales Setup – danach nicht mehr verwendet"],
        ["electroswap_admin",    "dbOwner",              "Restore, Schema-Änderungen"],
        ["electroswap_app",      "readWrite",            "Flask-Applikation (lesen + schreiben)"],
        ["electroswap_readonly", "read",                 "Backup-Prozess (nur lesen)"],
    ],
    col_widths=[4.5, 3.5, 8.5]
)

add_heading(tmp, "Berechtigungsmatrix – Datenbank", level=3)

add_table(tmp,
    ["Operation", "admin", "app (Flask)", "readonly"],
    [
        ["Lesen (find, aggregate)",           "✓", "✓", "✓"],
        ["Schreiben (insert, update, delete)", "✓", "✓", "✗"],
        ["Collections verwalten (drop)",       "✓", "✗", "✗"],
        ["Backup erstellen (mongodump)",        "✓", "✗", "✓"],
        ["Restore ausführen (--drop)",          "✓", "✗", "✗"],
        ["Benutzer verwalten",                 "✓", "✗", "✗"],
    ],
    col_widths=[6.5, 2.5, 3, 2.5]
)

add_para(tmp,
    "Die Benutzer werden automatisch beim ersten Container-Start durch das Skript "
    "mongo-init/init-users.js angelegt. Der Flask-App wird ausschliesslich der "
    "electroswap_app Benutzer (readWrite) übergeben – er hat keine Möglichkeit, "
    "die Datenbankstruktur zu verändern oder einen Restore durchzuführen."
)

add_heading(tmp, "Befehle: Benutzer anzeigen und verwalten", level=3)

add_para(tmp,
    "Die folgenden mongosh-Befehle zeigen die wichtigsten Verwaltungsoperationen. "
    "Sie werden innerhalb des Docker-Containers ausgeführt:"
)

add_code(tmp, [
    "# Verbindung als Admin herstellen",
    "docker exec -it electroswap_mongo mongosh electroswap \\",
    '  --username electroswap_admin --password "AdminPass123!" \\',
    "  --authenticationDatabase electroswap",
    "",
    "# Alle Datenbankbenutzer anzeigen",
    "db.getUsers()",
    "",
    "# Alle Applikationsbenutzer (customer/admin) anzeigen",
    'db.users.find({}, { username: 1, email: 1, role: 1, _id: 0 })',
    "",
    "# Benutzerrolle auf admin setzen",
    'db.users.updateOne({ username: "demo_customer" }, { $set: { role: "admin" } })',
    "",
    "# Benutzerrolle zurücksetzen",
    'db.users.updateOne({ username: "demo_customer" }, { $set: { role: "customer" } })',
])

tmp.add_page_break()

# ────────────────────────────────────────────────────────────────
# 3.2 – LIVE-DEMO ZUGRIFFSBERECHTIGUNGEN
# ────────────────────────────────────────────────────────────────
add_demo_banner(tmp, "3.2", "Zugriffsberechtigungen zeigen")

add_info_box(tmp,
    "Ziel: customer vs. admin im Browser zeigen, MongoDB-Benutzer mit db.getUsers() "
    "zeigen, Least-Privilege demonstrieren (readonly darf nicht schreiben), "
    "Rollenänderung live durchführen."
)

add_heading(tmp, "Vorbereitung (vor der Aufnahme)", level=3)

add_code(tmp, [
    "# Stack starten und Seed-Daten laden",
    "docker compose up --build -d",
    "docker exec electroswap_app1 python seed_data.py",
    "docker compose ps   # alle 5 Services müssen Up sein",
])

add_para(tmp, "Browser öffnen: http://localhost → Startseite muss laden.", italic=True)

add_step(tmp, 1, "Customer einloggen – Admin-Zugriff ist gesperrt",
    "Browser zeigen: customer@electroswap.ch hat keinen Admin-Bereich")
add_code(tmp, [
    "# Browser: http://localhost/auth/login",
    "#   E-Mail:   customer@electroswap.ch  |  Passwort: customer123",
    "#",
    "# Zeigen:",
    "#   → Kein Admin-Link im Navigationsmenü",
    "#   → http://localhost/admin/ direkt aufrufen",
    "#   → Flash-Meldung 'Access denied' erscheint, Weiterleitung zur Startseite",
])

add_step(tmp, 2, "Admin einloggen – voller Zugriff",
    "Browser zeigen: admin@electroswap.ch sieht das Admin-Dashboard")
add_code(tmp, [
    "# Ausloggen, dann neu einloggen:",
    "#   E-Mail:   admin@electroswap.ch  |  Passwort: admin123",
    "#",
    "# Zeigen:",
    "#   → Admin-Link im Navigationsmenü erscheint",
    "#   → /admin/ öffnet das Dashboard mit Statistiken",
    "#   → /admin/products/ zeigt Produkttabelle mit Bearbeiten-Buttons",
])

add_step(tmp, 3, "MongoDB-Datenbankbenutzer anzeigen",
    "Terminal: db.getUsers() zeigt drei Benutzer mit unterschiedlichen Rollen")
add_code(tmp, [
    "docker exec electroswap_mongo mongosh electroswap \\",
    '  --username electroswap_admin --password "AdminPass123!" \\',
    "  --authenticationDatabase electroswap \\",
    "  --eval 'db.getUsers()'",
    "",
    "# Ausgabe: electroswap_admin (dbOwner), electroswap_app (readWrite),",
    "#          electroswap_readonly (read)",
])

add_step(tmp, 4, "Least-Privilege: Readonly darf nicht schreiben",
    "Fehler zeigen: readonly bekommt 'not authorized' beim Schreiben")
add_code(tmp, [
    "docker exec electroswap_mongo mongosh electroswap \\",
    '  --username electroswap_readonly --password "ReadPass789!" \\',
    "  --authenticationDatabase electroswap \\",
    "  --eval 'db.users.insertOne({ test: \"unauthorized\" })'",
    "",
    "# Ausgabe: MongoServerError: not authorized on electroswap ...",
])

add_step(tmp, 5, "Live-Rollenänderung: customer → admin → customer",
    "Terminal + Browser: Rollenänderung wirkt sofort nach erneutem Login")
add_code(tmp, [
    "# Rolle auf admin setzen (PowerShell: $set → `$set)",
    "docker exec electroswap_mongo mongosh \\",
    '  "mongodb://electroswap_admin:AdminPass123%21@localhost:27017/electroswap?authSource=electroswap" \\',
    '  --eval "db.users.updateOne({username:\'demo_customer\'},{`$set:{role:\'admin\'}});',
    "           db.users.findOne({username:'demo_customer'},{username:1,role:1,_id:0})\"",
    "",
    "# Ausgabe: { username: 'demo_customer', role: 'admin' }",
    "# → Im Browser neu einloggen → Admin-Dashboard zugänglich",
    "",
    "# Danach zurücksetzen:",
    '  --eval "db.users.updateOne({username:\'demo_customer\'},{`$set:{role:\'customer\'}})"',
])

add_info_box(tmp,
    "Alternativ: bash demo_rollen.sh ausführen – führt alle 5 Schritte automatisch "
    "durch und pausiert zwischen den Blöcken für die Kamera.",
    fill="FFF8E1"
)

tmp.add_page_break()

# ────────────────────────────────────────────────────────────────
# 3.3 – BACKUP-KONZEPT
# ────────────────────────────────────────────────────────────────
add_heading(tmp, "3.3  Backup der Datenbank: Konzept", level=2)

add_para(tmp,
    "Backups schützen vor Datenverlust durch Hardware-Ausfall, Fehlbedienung oder "
    "Ransomware. ElectroSwap setzt auf das offizielle MongoDB-Tool mongodump, das "
    "alle Collections in ein komprimiertes Archiv exportiert, das mit mongorestore "
    "vollständig wiederhergestellt werden kann."
)

add_heading(tmp, "Strategie", level=3)

add_table(tmp,
    ["Aspekt", "Entscheidung", "Begründung"],
    [
        ["Tool",             "mongodump / mongorestore",
         "Offizielle MongoDB-Tools, BSON-Format, vollständig kompatibel"],
        ["Backup-Benutzer",  "electroswap_readonly",
         "Least-Privilege: Lesen genügt – kein Schreibzugriff nötig"],
        ["Restore-Benutzer", "electroswap_admin",
         "--drop erfordert Schreibrecht (dbOwner)"],
        ["Format",           "Komprimiertes Archiv (.gz)",
         "Platzsparend, portabel, einfach zu übertragen"],
        ["Dateiname",        "electroswap_YYYYMMDD_HHMMSS.gz",
         "Eindeutiger Zeitstempel, automatisch sortierbar"],
        ["Aufbewahrung",     "Automatische Bereinigung nach 7 Tagen",
         "Verhindert unkontrollierten Speicherverbrauch"],
    ],
    col_widths=[4, 5, 7.5]
)

add_heading(tmp, "Gesicherte Collections", level=3)

add_table(tmp,
    ["Collection", "Inhalt", "Warum wichtig?"],
    [
        ["users",     "Konten, Rollen, Passwort-Hashes, Adressen",
         "Grundlage für Login und Berechtigungen"],
        ["products",  "Produktkatalog mit Spezifikationen",
         "Kern des Shop-Angebots"],
        ["orders",    "Bestellhistorie mit Preis-Snapshot",
         "Rechtlich relevant – muss unveränderbar sein"],
        ["baskets",   "Aktive Warenkörbe",
         "Kundendaten – Verlust bedeutet Kaufabbruch"],
        ["wishlists", "Wunschlisten",
         "Kundenpräferenzen"],
        ["reviews",   "Produktbewertungen (verifiziert)",
         "Vertrauen und Produktqualität"],
    ],
    col_widths=[3, 5.5, 8]
)

add_heading(tmp, "Backup-Befehl (mit Authentifizierung)", level=3)

add_code(tmp, [
    "# Backup erstellen – readonly-User (Least Privilege)",
    "docker exec electroswap_mongo mongodump \\",
    '  --username electroswap_readonly --password "ReadPass789!" \\',
    "  --authenticationDatabase electroswap --db electroswap \\",
    "  --archive | gzip > ./backup/dumps/electroswap_$(date +%Y%m%d_%H%M%S).gz",
    "",
    "# Automatisches Skript (Zeitstempel + Bereinigung alter Backups):",
    "bash demo_backup.sh",
])

add_heading(tmp, "Restore-Befehl (mit Authentifizierung)", level=3)

add_code(tmp, [
    "# Restore aus Archiv – Admin-User erforderlich (--drop benötigt Schreibrecht)",
    "gunzip -c ./backup/dumps/electroswap_TIMESTAMP.gz \\",
    "  | docker exec -i electroswap_mongo mongorestore \\",
    '      --username electroswap_admin --password "AdminPass123!" \\',
    "      --authenticationDatabase electroswap --db electroswap \\",
    "      --drop --archive",
    "",
    "# --drop löscht bestehende Collections vor dem Wiedereinspielen,",
    "# um doppelte Dokumente zu verhindern.",
    "",
    "# Automatisches Skript (wählt neuestes Backup automatisch):",
    "bash demo_restore.sh",
])

add_heading(tmp, "Warum Readonly für Backup?", level=3)

add_para(tmp,
    "Das Least-Privilege-Prinzip besagt: Jeder Prozess erhält nur die minimal nötigen Rechte. "
    "Ein Backup liest Daten nur – er schreibt nichts. Wird das Backup-Skript kompromittiert, "
    "kann ein Angreifer mit dem Readonly-User keine Daten verändern oder löschen. "
    "Der Restore hingegen benötigt den Admin-User, weil --drop das Löschen bestehender "
    "Collections erfordert."
)

tmp.add_page_break()

# ────────────────────────────────────────────────────────────────
# 3.4 – LIVE-DEMO BACKUP & RESTORE
# ────────────────────────────────────────────────────────────────
add_demo_banner(tmp, "3.4", "Backup und Restore mit Authentifizierung")

add_info_box(tmp,
    "Ziel: (1) Falsches Passwort → Auth-Fehler, (2) Korrektes Backup mit readonly-User, "
    "(3) Datenverlust simulieren, (4) Restore mit Admin-User, (5) Daten zurück im Browser."
)

add_heading(tmp, "Vorbereitung (vor der Aufnahme)", level=3)

add_code(tmp, [
    "docker compose up -d && docker compose ps",
    "docker exec electroswap_app1 python seed_data.py",
])

add_step(tmp, 1, "Ausgangszustand der Datenbank zeigen",
    "Anzahl Dokumente vor Backup bestätigen")
add_code(tmp, [
    "docker exec electroswap_mongo mongosh electroswap \\",
    '  --username electroswap_admin --password "AdminPass123!" \\',
    "  --authenticationDatabase electroswap \\",
    "  --eval 'print(\"products: \" + db.products.countDocuments() +",
    '          \" | users: \" + db.users.countDocuments())',
    "# Ausgabe: products: 35 | users: 36",
])

add_step(tmp, 2, "Backup mit FALSCHEM Passwort – Fehler zeigen",
    "Authentifizierung schlägt fehl → Daten sind geschützt")
add_code(tmp, [
    "bash demo_backup_fail.sh",
    "",
    "# Ausgabe: Authentication failed. – MongoDB verweigert den Zugriff",
])

add_step(tmp, 3, "Backup mit KORREKTEM Passwort (readonly-User)",
    "Backup-Datei wird erstellt, Grösse sichtbar")
add_code(tmp, [
    "bash demo_backup.sh",
    "ls -lh backup/dumps/   # Datei electroswap_TIMESTAMP.gz ~12 KB",
])

add_step(tmp, 4, "Datenverlust simulieren – Produkte löschen",
    "Alle Produkte entfernen, Browser-Katalog ist leer")
add_code(tmp, [
    "docker exec electroswap_mongo mongosh electroswap \\",
    '  --username electroswap_admin --password "AdminPass123!" \\',
    "  --authenticationDatabase electroswap \\",
    "  --eval 'db.products.deleteMany({}); print(db.products.countDocuments())'",
    "# Ausgabe: 0  →  Browser http://localhost/products/ zeigt leeren Katalog",
])

add_step(tmp, 5, "Restore – Daten werden wiederhergestellt",
    "Skript zeigt automatisch Vorher/Nachher-Zustand")
add_code(tmp, [
    "bash demo_restore.sh",
    "# Ausgabe: products: 0 → Restore → products: 35",
    "",
    "# Browser: http://localhost/products/ → Katalog vollständig zurück",
])

add_step(tmp, 6, "Least-Privilege: Readonly darf nicht restoren",
    "Readonly-User bekommt 'not authorized' beim --drop Befehl")
add_code(tmp, [
    "docker exec electroswap_mongo mongorestore \\",
    '  --username electroswap_readonly --password "ReadPass789!" \\',
    "  --authenticationDatabase electroswap --db electroswap --drop \\",
    "  /dump/backup/electroswap",
    "# Ausgabe: not authorized on electroswap to execute command { drop: ... }",
])

tmp.add_page_break()

# ────────────────────────────────────────────────────────────────
# 3.5 – KONZEPT HORIZONTALE SKALIERUNG
# ────────────────────────────────────────────────────────────────
add_heading(tmp, "3.5  Konzept für horizontale Skalierung", level=2)

add_para(tmp,
    "Eine einzelne Flask-Instanz ist ein Single Point of Failure (SPOF): Fällt sie aus, "
    "ist die gesamte Applikation nicht mehr erreichbar. Ausserdem ist die Leistung auf "
    "einen einzigen Server begrenzt. Horizontale Skalierung bedeutet: Statt einen Server "
    "zu vergrössern (vertikal), werden mehrere identische Server-Instanzen parallel betrieben."
)

add_heading(tmp, "Architektur", level=3)

add_code(tmp, [
    "                     ┌────────────────┐",
    "       Internet  →   │     Nginx      │  Load Balancer (Port 80)",
    "                     │  Round-Robin   │",
    "                     └───────┬────────┘",
    "               ┌─────────────┼─────────────┐",
    "               ▼             ▼             ▼",
    "         ┌──────────┐  ┌──────────┐  ┌──────────┐",
    "         │  app1    │  │  app2    │  │  app3    │",
    "         │  Flask   │  │  Flask   │  │  Flask   │",
    "         └────┬─────┘  └────┬─────┘  └────┬─────┘",
    "              └─────────────┼─────────────┘",
    "                            ▼",
    "                    ┌───────────────┐",
    "                    │   MongoDB     │",
    "                    └───────────────┘",
])

add_heading(tmp, "Warum ist horizontale Skalierung möglich?", level=3)

add_para(tmp,
    "Eine Applikation kann nur dann horizontal skaliert werden, wenn sie zustandslos "
    "(stateless) ist – d.h. kein Node speichert lokalen Zustand. ElectroSwap erfüllt "
    "diese Bedingung vollständig:"
)

add_table(tmp,
    ["Aspekt", "Lösung", "Bedeutung für Skalierung"],
    [
        ["Sessions",    "Signierte Client-seitige Cookies (Flask SECRET_KEY)",
         "Kein Node muss sich den Benutzer merken – jeder Node kann den Request verarbeiten"],
        ["Datenpersistenz", "Alle Nodes teilen dieselbe MongoDB-Instanz",
         "Konsistente Daten ohne Synchronisation zwischen Nodes"],
        ["Dateisystem",     "Keine lokalen Uploads, Bilder via URL",
         "Kein geteilter Speicher nötig"],
    ],
    col_widths=[3.5, 6, 7]
)

add_heading(tmp, "Nginx als Load Balancer (Round-Robin)", level=3)

add_para(tmp,
    "Nginx verteilt eingehende Requests reihum auf die drei App-Nodes: "
    "Request 1 → app1, Request 2 → app2, Request 3 → app3, Request 4 → app1, … "
    "Fällt ein Node aus, erkennt Nginx dies und leitet automatisch nur noch an die "
    "gesunden Nodes weiter – ohne Konfigurationsänderung."
)

add_table(tmp,
    ["Komponente", "Technologie", "Funktion"],
    [
        ["Load Balancer", "Nginx 1.25",     "Verteilt Requests Round-Robin, überwacht Node-Gesundheit"],
        ["App-Nodes",     "Flask × 3",      "Verarbeiten Requests unabhängig, verbinden sich alle mit MongoDB"],
        ["Datenbank",     "MongoDB 7.0",    "Gemeinsamer, persistenter Datenspeicher für alle Nodes"],
    ],
    col_widths=[4, 4, 8.5]
)

add_heading(tmp, "Replikation und Sharding (Weiterführend)", level=3)

add_para(tmp,
    "In der Demo läuft MongoDB als Single Node. Für produktive Systeme gibt es zwei "
    "Erweiterungsstufen: Ein Replica Set spiegelt Daten auf mehrere Nodes – fällt der "
    "Primary-Node aus, wählen die Secondary-Nodes automatisch einen neuen Primary (Failover). "
    "Sharding teilt sehr grosse Datenmengen anhand eines Shard-Keys auf mehrere Server auf, "
    "sodass auch die Datenbank horizontal skaliert werden kann."
)

add_table(tmp,
    ["Einschränkung in dieser Demo", "Grund"],
    [
        ["MongoDB läuft als Single Node",    "Kein Replica Set → DB ist nicht redundant"],
        ["Alle Nodes auf demselben Host",    "Kein echter Hardware-Ausfall simulierbar"],
        ["Kein HTTPS",                       "Für Produktion wäre TLS-Termination bei Nginx nötig"],
    ],
    col_widths=[7, 9.5]
)

tmp.add_page_break()

# ────────────────────────────────────────────────────────────────
# 3.6 – LIVE-DEMO HORIZONTALE SKALIERUNG
# ────────────────────────────────────────────────────────────────
add_demo_banner(tmp, "3.6", "Horizontale Skalierung mit 3 Nodes realisieren")

add_info_box(tmp,
    "Ziel: Stack mit 3 App-Nodes starten, Load Balancing im Log sichtbar machen, "
    "Node-Ausfall simulieren, DB-Authentifizierung der Nodes zeigen."
)

add_heading(tmp, "Vorbereitung (vor der Aufnahme)", level=3)

add_code(tmp, [
    "docker compose down   # alten Stack beenden",
    "cd C:\\Projects\\ElectroSwap\\ElectroSwap",
])

add_step(tmp, 1, "Stack mit allen 5 Services starten",
    "Terminal: alle Services starten und Logs beobachten")
add_code(tmp, [
    "docker compose up --build",
    "# Im zweiten Terminal prüfen:",
    "docker compose ps",
    "# Erwartet: electroswap_nginx Up | app1 Up | app2 Up | app3 Up | mongo Up (healthy)",
])

add_step(tmp, 2, "Seed-Daten laden und App im Browser zeigen",
    "Browser: Produkte und Admin-Dashboard zeigen")
add_code(tmp, [
    "docker exec electroswap_app1 python seed_data.py",
    "# Browser: http://localhost → Startseite mit Produkten",
])

add_step(tmp, 3, "Load Balancing live zeigen",
    "Logs: Requests gehen reihum an app1, app2, app3")
add_code(tmp, [
    "# Terminal 1 – Logs beobachten:",
    "docker compose logs -f app1 app2 app3",
    "",
    "# Terminal 2 – 9 Requests senden:",
    'for i in $(seq 1 9); do curl -s -o /dev/null -w "%{http_code}\\n" http://localhost/; done',
    "",
    "# Im Log erkennbar: app1 / app2 / app3 erhalten je 3 Requests (Round-Robin)",
])

add_step(tmp, 4, "Node-Ausfall simulieren",
    "app2 stoppen – App läuft weiter über app1 und app3")
add_code(tmp, [
    "docker compose stop app2",
    "docker compose ps          # app2 ist Exited",
    "curl http://localhost/     # gibt 200 OK zurück (app1 oder app3 antwortet)",
    "docker compose start app2  # app2 wieder starten",
])

add_step(tmp, 5, "DB-Authentifizierung der App-Nodes zeigen",
    "Alle Nodes verbinden sich als electroswap_app (readWrite)")
add_code(tmp, [
    "docker exec electroswap_app1 env | grep MONGO_URI",
    "# Ausgabe: mongodb://electroswap_app:AppPass456!@mongo:27017/electroswap?authSource=electroswap",
    "",
    "# Verbindung direkt testen:",
    "docker exec electroswap_mongo mongosh electroswap \\",
    '  --username electroswap_app --password "AppPass456!" \\',
    "  --authenticationDatabase electroswap \\",
    "  --eval 'print(db.products.countDocuments())'",
])

add_step(tmp, 6, "Stack beenden",
    "Sauber stoppen")
add_code(tmp, [
    "docker compose down        # Stoppen, Volumes bleiben erhalten",
    "docker compose down -v     # Vollständiger Reset inkl. Volumes",
])

# ─── Kapitel-3-Inhalt in Hauptdokument einfügen ───────────────
# XML-Elemente aus tmp in doc einfügen (vor ref_elem)
tmp_body_elements = list(tmp.element.body)
for elem in tmp_body_elements:
    # Abschluss-sectPr überspringen
    if elem.tag.endswith('}sectPr'):
        continue
    new_elem = copy.deepcopy(elem)
    if ref_elem is not None:
        ref_elem.addprevious(new_elem)
    else:
        doc.element.body.append(new_elem)

# ─── Speichern ───────────────────────────────────────────────
doc.save(OUT)
print(f"Verbesserte Dokumentation gespeichert: {OUT}")
