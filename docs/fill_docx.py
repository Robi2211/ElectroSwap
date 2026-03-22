"""
Füllt die Abschnitte 3.2 – 3.6 im Word-Dokument mit dem richtigen Inhalt
im gleichen Stil wie 3.1.
"""

import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_IN  = "inf165-lb2-dokumentation-Ermel-Felber-2026  (1).docx"
DOC_OUT = "inf165-lb2-dokumentation-Ermel-Felber-2026  (1).docx"

doc = Document(DOC_IN)

XML_SP = "{http://www.w3.org/XML/1998/namespace}space"

# ──────────────────────────────────────────────────────────────────────────────
# Paragraph builders
# ──────────────────────────────────────────────────────────────────────────────

def _pStyle(val):
    el = OxmlElement("w:pStyle"); el.set(qn("w:val"), val); return el

def _run(text, bold=False, italic=False, mono=False, color=None):
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    if mono:
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), "Courier New"); rf.set(qn("w:hAnsi"), "Courier New")
        rPr.append(rf)
        for tag in ("w:sz", "w:szCs"):
            sz = OxmlElement(tag); sz.set(qn("w:val"), "18"); rPr.append(sz)
    if color:
        c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
    if len(rPr):
        r.append(rPr)
    t = OxmlElement("w:t"); t.set(XML_SP, "preserve"); t.text = text
    r.append(t)
    return r

def text_para(text, bold=False):
    """Normal body paragraph (style='Text')."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr"); pPr.append(_pStyle("Text")); p.append(pPr)
    p.append(_run(text, bold=bold))
    return p

def bold_label(label, rest=""):
    """Paragraph with bold label + normal continuation."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr"); pPr.append(_pStyle("Text")); p.append(pPr)
    p.append(_run(label, bold=True))
    if rest:
        p.append(_run(rest))
    return p

def bullet(text, numId="3"):
    """Bullet-list paragraph (Text style + numPr)."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr"); pPr.append(_pStyle("Text"))
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    nId  = OxmlElement("w:numId"); nId.set(qn("w:val"), numId)
    numPr.append(ilvl); numPr.append(nId); pPr.append(numPr)
    p.append(pPr); p.append(_run(text))
    return p

def code_para(line):
    """Single line of code: gray background, Courier New, blue left border."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr"); pPr.append(_pStyle("Normal"))

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2"); pPr.append(shd)

    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "6")
    left.set(qn("w:space"), "4"); left.set(qn("w:color"), "4472C4")
    pBdr.append(left); pPr.append(pBdr)

    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "360"); pPr.append(ind)
    p.append(pPr)

    stripped = line.strip()
    is_comment = stripped.startswith("#") or stripped.startswith("//") or stripped.startswith("→")
    color = "6A737D" if is_comment else None
    if line or stripped:
        p.append(_run(line, mono=True, color=color))
    return p

def code_block(code_str):
    """Convert multiline string into list of code_para elements."""
    lines = code_str.strip("\n").split("\n")
    return [code_para(ln) for ln in lines]

def make_table(headers, rows, col_widths=None):
    """Create a Table Grid table; first row is blue header."""
    tbl = OxmlElement("w:tbl")

    tblPr = OxmlElement("w:tblPr")
    ts = OxmlElement("w:tblStyle"); ts.set(qn("w:val"), "TableGrid"); tblPr.append(ts)
    tw = OxmlElement("w:tblW"); tw.set(qn("w:w"), "0"); tw.set(qn("w:type"), "auto")
    tblPr.append(tw)
    # No table indent
    tind = OxmlElement("w:tblInd"); tind.set(qn("w:w"), "0"); tind.set(qn("w:type"), "dxa")
    tblPr.append(tind)
    tbl.append(tblPr)

    tblGrid = OxmlElement("w:tblGrid")
    n = len(headers)
    width = 8748 // n  # approx full page width
    if col_widths:
        widths = col_widths
    else:
        widths = [width] * n
    for w_ in widths:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w_)); tblGrid.append(gc)
    tbl.append(tblGrid)

    def make_row(cells, is_header=False):
        tr = OxmlElement("w:tr")
        for i, cell_text in enumerate(cells):
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(widths[i] if i < len(widths) else width))
            tcW.set(qn("w:type"), "dxa"); tcPr.append(tcW)
            if is_header:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "4472C4"); tcPr.append(shd)
            tc.append(tcPr)

            # Support multi-line cell content (split by \n)
            lines = str(cell_text).split("\n")
            for li, line in enumerate(lines):
                cp = OxmlElement("w:p")
                cpPr = OxmlElement("w:pPr"); cpPr.append(_pStyle("Text")); cp.append(cpPr)
                cr = OxmlElement("w:r")
                crPr = OxmlElement("w:rPr")
                if is_header:
                    crPr.append(OxmlElement("w:b"))
                    col_el = OxmlElement("w:color"); col_el.set(qn("w:val"), "FFFFFF")
                    crPr.append(col_el)
                cr.append(crPr)
                ct = OxmlElement("w:t"); ct.set(XML_SP, "preserve"); ct.text = line
                cr.append(ct); cp.append(cr); tc.append(cp)
            tr.append(tc)
        return tr

    tbl.append(make_row(headers, is_header=True))
    for row in rows:
        tbl.append(make_row(row))
    return tbl

def empty_para():
    """Empty paragraph for spacing."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr"); pPr.append(_pStyle("Text")); p.append(pPr)
    return p

def h2_para(text):
    """Heading 2 paragraph."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr"); pPr.append(_pStyle("Heading2")); p.append(pPr)
    p.append(_run(text))
    return p

# ──────────────────────────────────────────────────────────────────────────────
# Insert helpers
# ──────────────────────────────────────────────────────────────────────────────

def insert_elements_after(ref_elem, elements):
    """Insert list of elements directly after ref_elem, preserving order."""
    cur = ref_elem
    for el in elements:
        cur.addnext(el)
        cur = el

def replace_paragraph_with_elements(placeholder_p, elements):
    """Replace a paragraph XML element with a list of elements."""
    parent = placeholder_p.getparent()
    idx = list(parent).index(placeholder_p)
    # Insert all new elements before placeholder
    cur = placeholder_p
    for el in reversed(elements):
        cur.addprevious(el)
    # Remove the placeholder
    parent.remove(placeholder_p)

# ──────────────────────────────────────────────────────────────────────────────
# Collect reference elements BEFORE any modification
# ──────────────────────────────────────────────────────────────────────────────

# Find placeholder paragraphs by text search (robust against index drift)
def find_para_elem(search_text):
    for p in doc.paragraphs:
        if search_text in p.text:
            return p._p
    return None

def find_para_elem_by_index(idx):
    return doc.paragraphs[idx]._p

p32_ref = find_para_elem("Führen Sie die Befehle auf zur Umsetzung in der Datenbank")
p33_ref = find_para_elem("Führen Sie den Befehl auf, um ein Backup")
p34_ref = find_para_elem("Führen Sie den Befehl auf, das Backup wiederherzustellen")
p35_ref = find_para_elem("Erstellen Sie ein Konzept, wie Sie Ihre Datenbank horizontal")
p_applikation_ref = find_para_elem("Technologie / Aufbau der Applikation")  # H2 anchor for 3.6

# Also need the Konzept heading for 3.5 to insert 3.6 after 3.5 content
p_applikation_h1 = None
for p in doc.paragraphs:
    if p.text.strip() == "Applikation" and p.style.name == "Heading 1":
        p_applikation_h1 = p._p
        break

# ──────────────────────────────────────────────────────────────────────────────
# 3.2 – Zugriffsberechtigung: Umsetzung
# ──────────────────────────────────────────────────────────────────────────────

els_32 = []

els_32.append(text_para("Applikationsebene – Decorator im Code", bold=True))
els_32.append(text_para(
    "Der @admin_required-Decorator in app/admin/routes.py schützt alle Admin-Routen:"
))
els_32 += code_block("""\
def admin_required(f):
    @wraps(f)
    @login_required
    def wrapper(*args, **kwargs):
        if not current_user.is_admin:
            flash("Access denied.", "error")
            return redirect(url_for("main.index"))
        return f(*args, **kwargs)
    return wrapper""")

els_32.append(empty_para())
els_32.append(text_para("Datenbankebene – MongoDB-Benutzer erstellen", bold=True))
els_32.append(text_para("Verbindung als Root-Benutzer:"))
els_32 += code_block("""\
mongosh -u root -p RootPass000! --authenticationDatabase admin""")

els_32.append(text_para("Datenbankbenutzer anlegen:"))
els_32 += code_block("""\
use electroswap

// Admin-Benutzer (Vollzugriff)
db.createUser({
  user: "electroswap_admin",
  pwd: "AdminPass123!",
  roles: [{ role: "dbOwner", db: "electroswap" }]
})

// App-Benutzer (Lesen + Schreiben)
db.createUser({
  user: "electroswap_app",
  pwd: "AppPass456!",
  roles: [{ role: "readWrite", db: "electroswap" }]
})

// Readonly-Benutzer (nur Lesen)
db.createUser({
  user: "electroswap_readonly",
  pwd: "ReadPass789!",
  roles: [{ role: "read", db: "electroswap" }]
})""")

els_32.append(text_para("Benutzer anzeigen und verwalten:"))
els_32 += code_block("""\
// Alle Benutzer anzeigen
db.getUsers()

// Einzelnen Benutzer anzeigen
db.getUser("electroswap_app")

// Passwort ändern
db.changeUserPassword("electroswap_app", "NeuesPasswort!")

// Benutzer löschen
db.dropUser("electroswap_readonly")""")

els_32.append(text_para("Als bestimmter Benutzer verbinden:"))
els_32 += code_block("""\
// App-Benutzer (readWrite)
mongosh "mongodb://electroswap_app:AppPass456!@localhost:27017/electroswap?authSource=electroswap"

// Readonly-Benutzer
mongosh "mongodb://electroswap_readonly:ReadPass789!@localhost:27017/electroswap?authSource=electroswap\"""")

els_32.append(text_para("Applikations-Rolle eines Benutzers verwalten:"))
els_32 += code_block("""\
// Benutzer auf Admin befördern
db.users.updateOne(
  { email: "max@example.com" },
  { $set: { role: "admin" } }
)

// Alle Admins anzeigen
db.users.find({ role: "admin" }, { username: 1, email: 1 })""")

# ──────────────────────────────────────────────────────────────────────────────
# 3.3 – Backup der DB
# ──────────────────────────────────────────────────────────────────────────────

els_33 = []

els_33.append(text_para(
    "Das Backup erfolgt mit dem offiziellen MongoDB-Tool mongodump. "
    "Als Backup-Benutzer wird electroswap_readonly verwendet "
    "(Least-Privilege-Prinzip: Backup benötigt nur Lesezugriff)."
))
els_33.append(empty_para())

els_33.append(text_para("Backup-Befehl mit Authentifizierung:", bold=True))
els_33 += code_block("""\
mongodump \\
  --host localhost:27017 \\
  --db electroswap \\
  --username electroswap_readonly \\
  --password ReadPass789! \\
  --authenticationDatabase electroswap \\
  --out ./backup/dumps/electroswap_$(date +%Y%m%d_%H%M%S)""")

els_33.append(empty_para())
els_33.append(text_para("Backup aus laufendem Docker-Container:", bold=True))
els_33 += code_block("""\
docker exec electroswap_mongo mongodump \\
  --db electroswap \\
  --username electroswap_readonly \\
  --password ReadPass789! \\
  --authenticationDatabase electroswap \\
  --archive | gzip > ./backup/dumps/electroswap_$(date +%Y%m%d_%H%M%S).gz""")

els_33.append(empty_para())
els_33.append(text_para("Backup-Skript:", bold=True))
els_33.append(text_para(
    "Das Skript backup/backup.sh erstellt automatisch einen Dump mit Zeitstempel, "
    "komprimiert diesen und bereinigt Backups älter als 7 Tage:"
))
els_33 += code_block("bash backup/backup.sh")

els_33.append(empty_para())
els_33.append(text_para("Gesicherte Collections:", bold=True))
els_33.append(make_table(
    ["Collection", "Inhalt"],
    [
        ["users",    "Benutzerkonten, Rollen, Adressen"],
        ["products", "Produktkatalog mit Spezifikationen"],
        ["baskets",  "Warenkörbe"],
        ["wishlists","Wunschlisten"],
        ["orders",   "Bestellhistorie"],
        ["reviews",  "Produktbewertungen"],
    ],
    col_widths=[2916, 5832]
))

# ──────────────────────────────────────────────────────────────────────────────
# 3.4 – Restore eines DB-Backups
# ──────────────────────────────────────────────────────────────────────────────

els_34 = []

els_34.append(text_para(
    "Der Restore erfolgt mit mongorestore. Als Restore-Benutzer wird "
    "electroswap_admin verwendet, da --drop Vollzugriff benötigt."
))
els_34.append(empty_para())

els_34.append(text_para("Restore-Befehl mit Authentifizierung:", bold=True))
els_34 += code_block("""\
mongorestore \\
  --host localhost:27017 \\
  --db electroswap \\
  --username electroswap_admin \\
  --password AdminPass123! \\
  --authenticationDatabase electroswap \\
  --drop \\
  ./backup/dumps/electroswap_20260320_120000/electroswap""")

els_34.append(text_para(
    "Hinweis: --drop löscht die bestehenden Collections vor dem Wiedereinspielen, "
    "damit keine doppelten Daten entstehen."
))
els_34.append(empty_para())

els_34.append(text_para("Restore aus komprimiertem Archiv (Docker):", bold=True))
els_34 += code_block("""\
gunzip -c ./backup/dumps/electroswap_20260320_120000.gz \\
  | docker exec -i electroswap_mongo mongorestore \\
      --db electroswap \\
      --username electroswap_admin \\
      --password AdminPass123! \\
      --authenticationDatabase electroswap \\
      --drop \\
      --archive""")

els_34.append(empty_para())
els_34.append(text_para("Restore-Skript:", bold=True))
els_34 += code_block("bash backup/restore.sh backup/dumps/electroswap_20260320_120000.tar.gz")

els_34.append(empty_para())
els_34.append(text_para("Berechtigungen beim Restore:", bold=True))
els_34.append(text_para(
    "Der Readonly-Benutzer darf keinen Restore durchführen – "
    "dies demonstriert das Least-Privilege-Prinzip:"
))
els_34 += code_block("""\
mongorestore \\
  --username electroswap_readonly \\
  --password ReadPass789! \\
  --authenticationDatabase electroswap \\
  --drop ./backup/dumps/electroswap/
// → Fehler: not authorized on electroswap to execute command { drop: ... }""")

# ──────────────────────────────────────────────────────────────────────────────
# 3.5 – Konzept für horizontale Skalierung
# ──────────────────────────────────────────────────────────────────────────────

els_35 = []

els_35.append(text_para(
    "Eine einzelne Flask-Instanz bildet einen Single Point of Failure und ist auf "
    "die Ressourcen eines einzigen Servers begrenzt. Bei erhöhter Last (z. B. "
    "Sale-Aktionen) soll die Applikation ohne Downtime horizontal skaliert werden."
))
els_35.append(empty_para())

els_35.append(text_para("Architektur:", bold=True))
els_35 += code_block("""\
                   ┌─────────────┐
     Internet  →   │    Nginx    │  Load Balancer (Port 80)
                   │  Round-Robin│
                   └──────┬──────┘
          ┌───────────────┼───────────────┐
          ▼               ▼               ▼
    ┌─────────┐     ┌─────────┐     ┌─────────┐
    │  app1   │     │  app2   │     │  app3   │
    │  :5000  │     │  :5000  │     │  :5000  │
    └────┬────┘     └────┬────┘     └────┬────┘
         └───────────────┼───────────────┘
                         ▼
                 ┌───────────────┐
                 │    MongoDB    │
                 │    :27017     │
                 └───────────────┘""")

els_35.append(empty_para())
els_35.append(text_para("Replikation:", bold=True))
els_35.append(text_para(
    "MongoDB Replica Sets spiegeln Daten auf mehrere Knoten. Ein Primary-Node nimmt "
    "alle Schreiboperationen entgegen; Secondary-Nodes replizieren die Daten "
    "automatisch. Fällt der Primary aus, wählen die Secondaries automatisch einen "
    "neuen Primary (Failover)."
))

els_35.append(empty_para())
els_35.append(text_para("Partitionierung (Sharding):", bold=True))
els_35.append(text_para(
    "Bei sehr grossen Datenmengen wird die Datenbank anhand eines Shard-Keys auf "
    "mehrere Shard-Nodes aufgeteilt. Jeder Shard enthält nur einen Teil der Daten. "
    "Ein Config-Server und Mongos-Router koordinieren die Anfragen."
))

els_35.append(empty_para())
els_35.append(text_para("Multimaster vs. Master-Slave:", bold=True))
els_35.append(make_table(
    ["Strategie", "Beschreibung", "Einsatz"],
    [
        ["Master-Slave",
         "Ein Primary schreibt, Secondary liest",
         "Standard MongoDB Replica Set (unsere Demo)"],
        ["Multimaster",
         "Mehrere Nodes können schreiben",
         "Geografisch verteilte Systeme"],
    ],
    col_widths=[2187, 3645, 2916]
))

els_35.append(empty_para())
els_35.append(text_para("Warum ist horizontale Skalierung möglich?", bold=True))
els_35.append(text_para(
    "ElectroSwap ist zustandslos (stateless) auf Applikationsebene:"
))
els_35.append(make_table(
    ["Aspekt", "Lösung"],
    [
        ["Sessions",
         "Signierte Client-Cookies (Flask SECRET_KEY) – keine serverseitige Speicherung nötig"],
        ["Persistenter Zustand",
         "Alle App-Nodes verbinden sich mit derselben MongoDB-Instanz"],
        ["Dateisystem",
         "Keine lokalen Datei-Uploads – kein geteilter Speicher nötig"],
    ],
    col_widths=[2916, 5832]
))

els_35.append(empty_para())
els_35.append(text_para("Nginx Load Balancer (Round-Robin):", bold=True))
els_35.append(text_para(
    "Jeder neue Request wird reihum an den nächsten App-Node weitergeleitet "
    "(nginx/nginx.conf):"
))
els_35 += code_block("""\
upstream electroswap_app {
    server app1:5000;
    server app2:5000;
    server app3:5000;
}

server {
    listen 80;
    location / {
        proxy_pass http://electroswap_app;
        proxy_set_header Host            $host;
        proxy_set_header X-Real-IP       $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
    }
}""")

els_35.append(empty_para())
els_35.append(text_para("Einschränkungen der Demo:", bold=True))
els_35.append(bullet("MongoDB läuft als Single Node (kein Replica Set) → Datenbankebene nicht redundant", "2"))
els_35.append(bullet("Alle Nodes laufen auf demselben Host → kein echter Hardware-Ausfall abgedeckt", "2"))
els_35.append(bullet("Kein HTTPS (TLS-Termination bei Nginx für Produktion notwendig)", "2"))

# ──────────────────────────────────────────────────────────────────────────────
# 3.6 – Horizontale Skalierung mit 3 Nodes realisieren
# ──────────────────────────────────────────────────────────────────────────────

els_36 = []

els_36.append(h2_para("Horizontale Skalierung mit 3 Nodes realisieren"))

els_36.append(text_para(
    "Die Infrastruktur wird mit Docker Compose umgesetzt: 3 Flask-App-Nodes, "
    "1 Nginx Load Balancer und 1 MongoDB-Instanz."
))
els_36.append(empty_para())

els_36.append(text_para("Voraussetzungen:", bold=True))
els_36 += code_block("""\
docker --version          # Docker 24+
docker compose version    # v2.20+""")

els_36.append(empty_para())
els_36.append(text_para("Infrastruktur starten:", bold=True))
els_36 += code_block("""\
# Im Projektverzeichnis
docker compose up --build

# Oder im Hintergrund
docker compose up --build -d""")

els_36.append(empty_para())
els_36.append(text_para("Laufende Services anzeigen:", bold=True))
els_36 += code_block("docker compose ps")
els_36.append(text_para("Erwartete Ausgabe:"))
els_36 += code_block("""\
NAME                  STATUS        PORTS
electroswap_nginx     Up            0.0.0.0:80->80/tcp
electroswap_app1      Up            5000/tcp
electroswap_app2      Up            5000/tcp
electroswap_app3      Up            5000/tcp
electroswap_mongo     Up (healthy)  0.0.0.0:27017->27017/tcp""")

els_36.append(empty_para())
els_36.append(text_para("Load Balancing verifizieren:", bold=True))
els_36 += code_block("""\
# 9 Requests senden – werden auf app1/app2/app3 verteilt
for i in $(seq 1 9); do curl -s -o /dev/null -w "%{http_code}\\n" http://localhost/; done

# Logs beobachten
docker compose logs -f app1 app2 app3""")

els_36.append(empty_para())
els_36.append(text_para("Node-Ausfall simulieren:", bold=True))
els_36 += code_block("""\
docker compose stop app2   # app2 ausfallen lassen
curl http://localhost/     # App bleibt erreichbar über app1 und app3
docker compose start app2  # app2 wieder hochfahren""")

els_36.append(empty_para())
els_36.append(text_para("Seed-Daten einspielen:", bold=True))
els_36 += code_block("docker exec electroswap_app1 python seed_data.py")

els_36.append(empty_para())
els_36.append(text_para("Datenbankverbindung prüfen:", bold=True))
els_36 += code_block("""\
# Verbindungsstring anzeigen
docker exec electroswap_app1 env | grep MONGO_URI
# → mongodb://electroswap_app:AppPass456!@mongo:27017/electroswap?authSource=electroswap""")

els_36.append(empty_para())
els_36.append(text_para("Infrastruktur beenden:", bold=True))
els_36 += code_block("""\
docker compose down      # Stoppen, Daten bleiben erhalten
docker compose down -v   # Kompletter Reset inkl. Volumes""")

# ──────────────────────────────────────────────────────────────────────────────
# Apply all changes (IMPORTANT: apply in document order)
# ──────────────────────────────────────────────────────────────────────────────

# 3.2 – replace placeholder
replace_paragraph_with_elements(p32_ref, els_32)

# 3.3 – replace placeholder
replace_paragraph_with_elements(p33_ref, els_33)

# 3.4 – replace placeholder
replace_paragraph_with_elements(p34_ref, els_34)

# 3.5 – replace placeholder
replace_paragraph_with_elements(p35_ref, els_35)

# 3.6 – insert before "Applikation" heading
for el in reversed(els_36):
    p_applikation_h1.addprevious(el)

# ──────────────────────────────────────────────────────────────────────────────
doc.save(DOC_OUT)
print("Gespeichert:", DOC_OUT)
